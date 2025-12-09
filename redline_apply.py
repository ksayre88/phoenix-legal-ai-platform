# phoenix_contracts/redline_apply.py

import io
from typing import List, Dict, Any, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------
# 1. Utilities for Word Redline XML Nodes
# ---------------------------------------------------------

def _make_insert_run(text: str):
    """
    Create a <w:ins> node containing the inserted text.
    """
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "Phoenix AI")
    ins.set(qn("w:date"), "2025-01-01T00:00:00Z")

    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    ins.append(run)
    return ins


def _make_delete_run(text: str):
    """
    Create a <w:del> node representing deleted text.
    """
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), "2")
    dele.set(qn("w:author"), "Phoenix AI")
    dele.set(qn("w:date"), "2025-01-01T00:00:00Z")

    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    dele.append(run)
    return dele


def _make_comment_ref(comment_id: int):
    """
    Insert a comment reference into a paragraph.
    """
    comment = OxmlElement("w:commentRangeStart")
    comment.set(qn("w:id"), str(comment_id))
    return comment


# ---------------------------------------------------------
# 2. Apply persona deltas to a paragraph
# ---------------------------------------------------------

def apply_deltas_to_paragraph(paragraph, delta: Dict[str, Any]):
    """
    paragraph: python-docx paragraph object
    delta: {insertions, deletions, replacements, comments}
    """

    p = paragraph._p  # underlying XML element

    # -------------------------------------
    # Insertions
    # -------------------------------------
    for ins_text in delta.get("insertions", []):
        ins_node = _make_insert_run(ins_text)
        p.append(ins_node)

    # -------------------------------------
    # Deletions
    # -------------------------------------
    for del_text in delta.get("deletions", []):
        del_node = _make_delete_run(del_text)
        p.append(del_node)

    # -------------------------------------
    # Replacements
    # -------------------------------------
    for rep in delta.get("replacements", []):
        old = rep.get("from")
        new = rep.get("to")
        if old and new:
            # Represent replacement as deletion + insertion
            del_node = _make_delete_run(old)
            ins_node = _make_insert_run(new)
            p.append(del_node)
            p.append(ins_node)

    # Comments are handled globally at the end (optional)
    return paragraph


# ---------------------------------------------------------
# 3. Main Redline Application Function
# ---------------------------------------------------------

def apply_redlines_to_docx(
    original_doc_bytes: bytes,
    deltas: List[Dict[str, Any]],
) -> bytes:
    """
    Takes:
      - original DOCX bytes
      - list of persona deltas (Part 2)
    Returns:
      - new DOCX bytes with redlines applied
    """

    doc = Document(io.BytesIO(original_doc_bytes))

    # Map CP hash → paragraph index
    cp_map = {}
    for i, p in enumerate(doc.paragraphs):
        key = p.text.strip()
        if key:
            import hashlib
            h = hashlib.sha256(key.encode("utf-8")).hexdigest()[:16]
            cp_map[h] = i

    # -------------------------------------
    # Apply deltas paragraph-by-paragraph
    # -------------------------------------
    for entry in deltas:
        cp_hash = entry.get("cp_hash")
        if not cp_hash:
            continue

        if cp_hash not in cp_map:
            continue  # no match found in DOCX, skip

        idx = cp_map[cp_hash]
        para = doc.paragraphs[idx]
        apply_deltas_to_paragraph(para, entry.get("delta", {}))

    # -------------------------------------
    # Save new DOCX
    # -------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
