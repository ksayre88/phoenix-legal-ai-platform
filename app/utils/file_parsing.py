import io
import re
from fastapi import HTTPException
from bs4 import BeautifulSoup

# Optional Imports handling
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    Traverses paragraphs and tables to get full content.
    """
    if Document is None: 
        raise HTTPException(status_code=500, detail="python-docx library not installed")
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip(): full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip(): full_text.append(para.text)
        return "\n".join(full_text).strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {e}")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from a PDF file using pypdf.
    """
    if PdfReader is None: return ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() for page in reader.pages if page.extract_text()]
        return "\n".join(pages)
    except Exception:
        return ""

def preprocess_document_from_upload(filename: str, file_bytes: bytes) -> dict:
    """
    Main entry point for file parsing. 
    Routes to specific handlers based on extension.
    """
    fn = filename.lower()
    text = ""
    
    try:
        if fn.endswith(".pdf"):
            if PdfReader:
                text = extract_text_from_pdf(file_bytes)
            else:
                return {"clean_text": "Error: pypdf library not installed.", "paragraphs": []}
                
        elif fn.endswith(".docx"):
            if Document:
                try:
                    text = extract_docx_text(file_bytes)
                except HTTPException:
                    return {"clean_text": "Error: Failed to parse DOCX structure.", "paragraphs": []}
            else:
                return {"clean_text": "Error: python-docx library not installed.", "paragraphs": []}
                
        elif fn.endswith(".html") or fn.endswith(".htm"):
            # RESTORED: BeautifulSoup logic matches original app.py
            soup = BeautifulSoup(file_bytes, "html.parser")
            text = soup.get_text(separator="\n")
            
        else:
            # Default to UTF-8 decoding for .txt or other files
            text = file_bytes.decode("utf-8", errors="ignore")
            
    except Exception as e:
        return {"clean_text": f"Error parsing document: {str(e)}", "paragraphs": []}
    
    if not text or not text.strip():
        return {"clean_text": "Error: File content was empty or unreadable.", "paragraphs": []}

    # Preprocess Policy Logic (matches original preprocess_policy_v4)
    # Filters out very short lines to reduce noise
    lines = [l.strip() for l in text.split("\n") if len(l.strip()) > 25]
    
    return {"clean_text": "\n".join(lines), "paragraphs": lines}