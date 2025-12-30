import io
import difflib
import re
from typing import List, Dict, Any, Tuple, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

# ---------------------------------------------------------
# 1. XML Helpers for Track Changes
# ---------------------------------------------------------

def _make_insert_run(text: str):
    """Creates a <w:ins> node (Track Changes Insertion)."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "0")
    ins.set(qn("w:author"), "Phoenix AI")
    ins.set(qn("w:date"), "2025-01-01T00:00:00Z")

    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    if text.strip():
        t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    ins.append(run)
    return ins

def _make_delete_run(text: str):
    """Creates a <w:del> node (Track Changes Deletion)."""
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), "0")
    dele.set(qn("w:author"), "Phoenix AI")
    dele.set(qn("w:date"), "2025-01-01T00:00:00Z")

    run = OxmlElement("w:r")
    
    # Visual strikethrough inside the deletion (double visual cue)
    rPr = OxmlElement("w:rPr")
    strike = OxmlElement("w:strike")
    rPr.append(strike)
    run.append(rPr)

    t = OxmlElement("w:t")
    if text.strip():
        t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    dele.append(run)
    return dele

# ---------------------------------------------------------
# 2. Fuzzy Matching Logic
# ---------------------------------------------------------

def find_fuzzy_match(paragraph_text: str, search_text: str, threshold=0.8) -> Tuple[int, int]:
    """
    Locates 'search_text' within 'paragraph_text' allowing for minor differences.
    Returns (start_index, end_index) or (-1, -1) if not found.
    """
    # 1. Clean both strings for comparison (ignore whitespace differences)
    # We use a SequenceMatcher to find the longest approximate match
    matcher = difflib.SequenceMatcher(None, paragraph_text, search_text, autojunk=False)
    
    match = matcher.find_longest_match(0, len(paragraph_text), 0, len(search_text))
    
    # Check if the match is good enough
    if match.size == 0:
        return -1, -1
        
    # Calculate similarity ratio of the matched block vs the search text
    matched_substring = paragraph_text[match.a : match.a + match.size]
    similarity = difflib.SequenceMatcher(None, matched_substring, search_text).ratio()
    
    if similarity < threshold:
        return -1, -1
        
    return match.a, match.a + match.size

# ---------------------------------------------------------
# 3. Paragraph Rebuilder
# ---------------------------------------------------------

def apply_deltas_to_paragraph(paragraph, delta: Dict[str, Any]):
    """
    Applies redlines using fuzzy matching and XML rebuilding.
    """
    full_text = paragraph.text
    replacements = delta.get("replacements", [])
    comments = delta.get("comments", [])

    if not replacements and not comments:
        return

    # Work on a copy of text to track indices if we had multiple replacements,
    # but for MVP we handle the *first valid* replacement to avoid overlap complexity.
    
    applied_change = False
    
    for rep in replacements:
        if applied_change: break # Only do one change per paragraph to prevent corruption
        
        old_str = rep.get("from", "").strip()
        new_str = rep.get("to", "").strip()
        
        if not old_str: continue

        # --- Try Exact Match First ---
        start_idx = full_text.find(old_str)
        end_idx = start_idx + len(old_str) if start_idx != -1 else -1

        # --- Try Fuzzy Match if Exact Fails ---
        if start_idx == -1:
            start_idx, end_idx = find_fuzzy_match(full_text, old_str)

        if start_idx != -1:
            # We found the text! Now split and rebuild.
            prefix = full_text[:start_idx]
            actual_old_text = full_text[start_idx:end_idx] # The text actually in the doc
            suffix = full_text[end_idx:]
            
            # 1. Clear the paragraph XML
            p_element = paragraph._p
            # Remove all content children (runs, ins, del) but keep properties
            for child in list(p_element):
                if child.tag.endswith("pPr"): continue
                p_element.remove(child)

            # 2. Rebuild Prefix (Normal)
            if prefix:
                r = p_element.add_r()
                t = OxmlElement("w:t")
                if prefix.strip(): t.set(qn("xml:space"), "preserve")
                t.text = prefix
                r.append(t)

            # 3. Add Deletion (The text we found in the doc)
            del_node = _make_delete_run(actual_old_text)
            p_element.append(del_node)

            # 4. Add Insertion (The AI's suggested text)
            if new_str:
                ins_node = _make_insert_run(new_str)
                p_element.append(ins_node)

            # 5. Rebuild Suffix (Normal)
            if suffix:
                r = p_element.add_r()
                t = OxmlElement("w:t")
                if suffix.strip(): t.set(qn("xml:space"), "preserve")
                t.text = suffix
                r.append(t)
                
            applied_change = True

    # --- Append Comments ---
    if comments:
        run = paragraph.add_run()
        run.add_break() 
        comment_text = "[AI: " + "; ".join(comments) + "]"
        run.text = comment_text
        run.bold = True
        run.font.color.rgb = RGBColor(0, 50, 150)

# ---------------------------------------------------------
# 4. Main Entry Point
# ---------------------------------------------------------

def apply_redlines_to_docx(
    original_doc_bytes: bytes,
    redlines: List[Dict[str, Any]],
) -> bytes:
    doc = Document(io.BytesIO(original_doc_bytes))

    # Optimization: Map paragraphs by text for faster lookup
    # We use a list of (index, text) to handle duplicate paragraphs correctly
    para_map = []
    for i, p in enumerate(doc.paragraphs):
        para_map.append({
            "index": i, 
            "text": p.text.strip(), 
            "obj": p
        })

    for entry in redlines:
        # The AI result structure (Service Layer)
        original_text = entry.get("original_text", "").strip()
        delta = entry.get("delta", {})
        
        if not original_text: continue

        # 1. Find the target paragraph
        # We look for the paragraph that contains the 'original_text' 
        # (or at least a significant part of it, in case of Stitching)
        
        target_para = None
        best_score = 0.0

        for p_entry in para_map:
            # Exact match check
            if p_entry["text"] == original_text:
                target_para = p_entry["obj"]
                break
            
            # Substring check (if the original_text was a stitched chunk, 
            # the paragraph might be just the body)
            if len(p_entry["text"]) > 50 and p_entry["text"] in original_text:
                target_para = p_entry["obj"]
                break
                
            # Fuzzy Similarity check (last resort)
            # This is slow, so we only do it if the paragraph is roughly the same length
            if abs(len(p_entry["text"]) - len(original_text)) < 50:
                score = difflib.SequenceMatcher(None, p_entry["text"], original_text).ratio()
                if score > 0.85 and score > best_score:
                    best_score = score
                    target_para = p_entry["obj"]

        # 2. Apply
        if target_para:
            apply_deltas_to_paragraph(target_para, delta)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()