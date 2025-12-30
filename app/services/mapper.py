import asyncio
import json
import re
import io
import base64
import gc  # Added for explicit garbage collection
from typing import List, Dict, Any, Optional
from pydantic import BaseModel, Field

# Docx dependencies for export
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

from app.utils.llm_client import call_ollama_generate
from app.core.config import settings

# ---------------------------------------------------------
# LAYER 0: CONFIGURATION
# ---------------------------------------------------------
GENERIC_TERMS = {
    "data", "information", "info", "details", "content", "services", 
    "user data", "personal data", "personal information", "pii", 
    "sensitive data", "consumer info", "material", "records", 
    "usage data", "electronic information", "other information",
    "third party data", "categories", "purposes", "identifying info",
    "commercial information", "information you provide", "information you entered"
}

BAD_ENTITIES = {
    "affiliates", "partners", "providers", "vendors", "third parties", 
    "subsidiaries", "contractors", "dealers", "agents", "service providers",
    "business partners", "advertising networks"
}

CORE_PI_FIELDS = ["Name", "Email Address", "Telephone Number", "Address"]

TERM_MAPPING = {
    "cell phone": "Telephone Number",
    "mobile number": "Telephone Number",
    "phone number": "Telephone Number",
    "phone": "Telephone Number",
    "email": "Email Address",
    "mail": "Email Address",
    "ip": "Ip Address",
    "internet protocol address": "Ip Address",
    "lat/long": "Geolocation Information",
    "gps": "Geolocation Information",
    "location": "Geolocation Information",
    "zip": "Address",
    "postal code": "Address",
    "device info": "Device Id",
    "hardware id": "Device Id",
    "advertising id": "Device Id",
    "idfa": "Device Id",
    "web beacon": "Web Beacons",
    "pixel": "Web Beacons"
}

# ---------------------------------------------------------
# DATA MODELS
# ---------------------------------------------------------
class DataFlowFinding(BaseModel):
    data_type: str = Field(..., description="The specific data element")
    action: str = Field(..., description="Collection, Sharing, Disclosure") 
    recipient: Optional[str] = Field(None, description="Specific Entity Name (e.g. Google, Meta)")

class ChunkResult(BaseModel):
    findings: List[DataFlowFinding] = []

class ProcessingPacket(BaseModel):
    chunk_id: int
    text: str
    definitions_context: str
    raw_llm_response: Optional[str] = None
    raw_findings: List[DataFlowFinding] = []
    verified_findings: List[DataFlowFinding] = []
    errors: List[str] = []

# ---------------------------------------------------------
# UTILS
# ---------------------------------------------------------
def clean_json_string(s: str) -> str:
    s = s.strip()
    if s.startswith("```json"): s = s[7:]
    if s.startswith("```"): s = s[3:]
    if s.endswith("```"): s = s[:-3]
    return s.strip()

def validate_llm_json(raw_output: str) -> List[DataFlowFinding]:
    clean = clean_json_string(raw_output)
    try:
        data = json.loads(clean)
        if isinstance(data, list):
             return [DataFlowFinding(**item) for item in data]
        result = ChunkResult(**data)
        return result.findings
    except:
        match = re.search(r"\{.*\}", clean, re.DOTALL)
        if match:
            try:
                data = json.loads(match.group(0))
                if "findings" in data:
                    return ChunkResult(**data).findings
                return [DataFlowFinding(**item) for item in data if isinstance(data, list)]
            except: pass
        return []

# ---------------------------------------------------------
# LAYER 1: STRICT GROUNDING (EVIDENCE CHECK)
# ---------------------------------------------------------
def evidence_check(findings: List[DataFlowFinding], text: str) -> List[DataFlowFinding]:
    valid_findings = []
    text_lower = text.lower()
    
    for f in findings:
        raw_dt = f.data_type.strip()
        if len(raw_dt) < 2: continue
        
        candidates = []
        candidates.append(raw_dt)
        clean_parens = re.sub(r'\s*\(.*?\)', '', raw_dt).strip()
        if clean_parens and clean_parens != raw_dt:
            candidates.append(clean_parens)
        if "/" in raw_dt:
            candidates.extend([x.strip() for x in raw_dt.split("/") if len(x.strip()) > 2])

        confirmed_candidates = []
        for cand in candidates:
            cand_lower = cand.lower()
            if len(cand) < 4:
                if re.search(r'\b' + re.escape(cand_lower) + r'\b', text_lower):
                    confirmed_candidates.append(cand)
            else:
                if cand_lower in text_lower:
                    confirmed_candidates.append(cand)
        
        if confirmed_candidates:
            # Pick LONGEST match to preserve specificity
            best_match = max(confirmed_candidates, key=len)
            f.data_type = best_match.title()
            valid_findings.append(f)
            
    return valid_findings

# ---------------------------------------------------------
# CORE ANALYSIS PIPELINE
# ---------------------------------------------------------

async def extract_definitions(text: str) -> str:
    intro_text = text[:8000] 
    prompt = (
        "Extract definitions from the text below.\n"
        "Summarize what specific data types are included in broad terms like 'Personal Information', 'Personal Data', or 'Usage Data'.\n"
        "Return a concise summary list.\n"
        f"TEXT:\n{intro_text}"
    )
    try:
        return await call_ollama_generate(settings.DEFAULT_MODEL_NAME, prompt, False, 256)
    except:
        return ""

def chunk_text(text: str, max_tokens=1000, overlap=100) -> List[str]:
    # OPTIMIZATION: 1000 tokens for fewer API calls and better context
    words = text.split()
    chunks = []
    step = max(1, max_tokens - overlap)
    for i in range(0, len(words), step):
        chunk = " ".join(words[i:i+max_tokens])
        chunks.append(chunk)
    return chunks

async def classify_chunks_parallel(
    chunks: List[str], 
    definitions_context: str, 
    prompt_builder=None, 
    verification_builder=None
) -> List[Dict[str, Any]]:
    
    # 10 Concurrent streams to saturate dual GPUs
    sem = asyncio.Semaphore(10)
    model = settings.DEFAULT_MODEL_NAME

    async def process_packet(text: str, idx: int):
        packet = ProcessingPacket(chunk_id=idx, text=text, definitions_context=definitions_context)
        
        async with sem:
            if prompt_builder: 
                prompt = prompt_builder(text, definitions_context)
            else: 
                prompt = (
                    f"Analyze this privacy policy text. Identify ALL data elements being collected or shared.\n"
                    f"If shared, identify the RECIPIENT NAME (e.g. 'Google', 'Dealers', 'Service Providers').\n"
                    f"Definitions Context: {definitions_context}\n\n"
                    f"TEXT:\n{text}"
                )
            
            try:
                # Reduced prediction length slightly to speed up turnaround
                resp_1 = await call_ollama_generate(model=model, prompt=prompt, json_mode=True, num_predict=512)
                packet.raw_llm_response = resp_1
                packet.raw_findings = validate_llm_json(resp_1)
                packet.verified_findings = evidence_check(packet.raw_findings, text)
            except Exception as e:
                packet.errors.append(f"Processing Error: {str(e)}")

            # --- GARBAGE COLLECTION 1: IMMEDIATE CLEANUP ---
            # Clear heavy fields immediately after processing to free memory
            packet.raw_llm_response = None
            packet.definitions_context = "" 
            
            return {
                "text": packet.text,
                "findings": [f.model_dump() for f in packet.verified_findings],
                "errors": packet.errors
            }

    tasks = [process_packet(chunks[i], i) for i in range(len(chunks))]
    results = await asyncio.gather(*tasks)
    
    # --- GARBAGE COLLECTION 2: FORCED FLUSH ---
    # Force a garbage collect after the heavy gather operation
    gc.collect()
    
    return results

# ---------------------------------------------------------
# LAYER 2-5: POST-PROCESSING GRAPH CONSTRUCTION
# ---------------------------------------------------------

def detect_company(text: str) -> str:
    """
    Detects the controller name using structural legal patterns.
    Updated with 'Section 10 / Contact Us' logic for policies like X/Twitter
    that hide the legal entity name at the very bottom.
    """
    # 1. Scan the END of the document (Footer/Contact Section)
    # This is often where "X Corp." or "Google LLC" is formally defined.
    outro = text[-5000:] 
    
    # Pattern: "data controller responsible... is [Entity]"
    match_controller = re.search(r"data controller responsible.*?is\s+([A-Z][a-zA-Z0-9\s\.,&]{2,50}?)(?:,|\.|with an address)", outro, re.IGNORECASE | re.DOTALL)
    if match_controller:
        return match_controller.group(1).strip()

    # Pattern: "Attn: ... [Entity] ... Address" or "[Entity] ... Attn:"
    # Matches: "X Corp.\nAttn: Privacy Policy Inquiry"
    match_attn = re.search(r"([A-Z][a-zA-Z0-9\s\.,&]{2,50}?)\s*\n\s*Attn:", outro, re.MULTILINE)
    if match_attn:
        candidate = match_attn.group(1).strip()
        if len(candidate) > 2 and "Service" not in candidate:
            return candidate

    # 2. Scan the Intro (Standard Header Logic)
    intro = text[:5000]
    
    # "We, [Entity], ..." (Friendly Legal)
    match_appositive = re.search(r"\bwe,\s+([A-Z][a-zA-Z0-9\s\.,&]{2,50}?)(?:,|process|provide)", intro)
    if match_appositive:
        candidate = match_appositive.group(1).strip().rstrip(".,")
        if " " in candidate: return candidate

    # "Entity respects..." (Classic)
    match_intro = re.search(r"^([A-Z][a-zA-Z0-9\s\.,&]{2,50}?)\s*(?:\(|LLC|Inc|Ltd).*?respects your privacy", intro, re.MULTILINE)
    if match_intro: return match_intro.group(1).strip()

    # "Provided/Operated by..."
    match_prov = re.search(r"(?:[Pp]rovided|[Cc]ontrolled|[Oo]perated)\s+by\s+([A-Z0-9][a-zA-Z0-9\s\.,&]{2,60}?)", intro)
    if match_prov: return match_prov.group(1).strip().rstrip(".,")

    # Legal Suffix Heuristic in Intro
    match_legal = re.search(r"([A-Z][a-zA-Z0-9\s&]{2,40}?)\s+(?:LLC|L\.L\.C\.|Inc\.?|Incorporated|Ltd\.?|Limited|Corp\.?|Corporation|GmbH|S\.A\.)", intro)
    if match_legal:
        candidate = match_legal.group(1).strip()
        if candidate.lower() not in ["the", "our", "this", "a", "statutory", "contact"]: return candidate

    # 3. Copyright Fallback
    match_copy = re.search(r"(?:©|Copyright)\s*(?:©|\d{4}|20\d{2})?-?(?:20\d{2})?,?\s+([A-Z][a-zA-Z0-9\s\.,&]{2,50})", outro, re.IGNORECASE)
    if match_copy:
        candidate = match_copy.group(1).strip()
        candidate = re.sub(r"(?i)\.?\s*all rights reserved.*", "", candidate).strip()
        if len(candidate) > 2 and len(candidate) < 60: return candidate

    return "Unknown Organization"

def extract_flows(classified_chunks: List[Dict[str, Any]], main_controller: str) -> List[Dict[str, Any]]:
    connections = {} 

    def map_recipient(raw_recip: Any) -> str:
        r = str(raw_recip or "").strip()
        r_lower = r.lower()
        
        # 1. Handle Self-References
        self_refs = {"we", "us", "our", "ours", "controller", "company", main_controller.lower()}
        if r_lower in self_refs: return "controller"
        if main_controller.lower() in r_lower and len(main_controller) > 3: return "controller"

        # 2. Handle Empty/Null
        if not r or r_lower in ["none", "null", "n/a", "unknown"]: return "unknown"

        # 3. Intelligent Bucketing
        if "partner" in r_lower: return "Third Party – Business Partners"
        if "dealer" in r_lower: return "Third Party – Dealerships"
        # Only flag Manufacturer if context suggests Auto OEM
        if ("manufacturer" in r_lower or "oem" in r_lower) and "device" not in r_lower: 
            return "Third Party – Manufacturers"

        if any(x in r_lower for x in ["social", "facebook", "meta", "twitter", "linkedin"]): return "Third Party – Social Media"
        if any(x in r_lower for x in ["advert", "marketing", "promo", "ad network"]): return "Third Party – Advertising"
        if any(x in r_lower for x in ["analytic", "track", "metric", "stat", "google"]): return "Processor – Analytics"
        if any(x in r_lower for x in ["gov", "law", "court", "police", "legal"]): return "Third Party – Legal Disclosure"
        
        return f"Processor – {r.title()}"

    for chunk_data in classified_chunks:
        findings = chunk_data.get("findings", [])
        for item in findings:
            d_type = item.get("data_type", "").strip()
            d_lower = d_type.lower()
            
            # --- LAYER 2: CANONICALIZATION ---
            if d_lower in TERM_MAPPING:
                d_type = TERM_MAPPING[d_lower]
                d_lower = d_type.lower()
            
            # --- LAYER 3: CONTEXTUAL EXPANSION ---
            types_to_process = [d_type]
            expansion_triggers = [
                "personal information", "personal data", "information you provide", 
                "information you entered", "commercial information", "account information"
            ]
            
            if d_lower in expansion_triggers:
                types_to_process = CORE_PI_FIELDS
            
            for final_dtype in types_to_process:
                if not final_dtype or len(final_dtype) < 2: continue
                
                # Filter generics ONLY if they weren't expanded/mapped
                if final_dtype.lower() in GENERIC_TERMS and final_dtype not in CORE_PI_FIELDS: continue
                if any(bad in final_dtype.lower() for bad in BAD_ENTITIES): continue
                
                # --- LAYER 4: ACTION ENFORCEMENT ---
                raw_recip = item.get("recipient", "")
                target_node = map_recipient(raw_recip)

                # Contextual Fallback
                if target_node == "unknown":
                     act = item.get("action", "Collection").lower()
                     if "shar" in act or "disclos" in act or "transfer" in act:
                          target_node = "Processor – Service Provider"
                     else:
                          target_node = "controller"

                if target_node == "controller":
                    final_action = "Collection"
                    tgt = "controller"
                    src = "user"
                else:
                    final_action = "Sharing"
                    tgt = target_node
                    src = "controller"
                
                key = (src, tgt, final_action)
                if key not in connections: connections[key] = set()
                connections[key].add(final_dtype)

    # --- SAFETY NET ---
    # Build text corpus for safety net scan
    full_text_corpus = " ".join([c.get("text", "").lower() for c in classified_chunks])
    
    # --- GARBAGE COLLECTION 3: DROP TEXT ---
    # Now that we've used the corpus, we can modify the input list to drop the heavy text
    for c in classified_chunks:
        c["text"] = "" 
    
    # Run Safety Net Logic
    if "share" in full_text_corpus or "disclose" in full_text_corpus:
        if "dealership" in full_text_corpus or "dealer" in full_text_corpus:
            dealer_key = ("controller", "Third Party – Dealerships", "Sharing")
            if dealer_key not in connections: connections[dealer_key] = set()
            for field in ["Name", "Email Address", "Telephone Number"]:
                connections[dealer_key].add(field)

    if "manufacturer" in full_text_corpus or "oem" in full_text_corpus:
         if "share" in full_text_corpus or "disclose" in full_text_corpus:
             mfg_key = ("controller", "Third Party – Manufacturers", "Sharing")
             if mfg_key not in connections: connections[mfg_key] = set()
             connections[mfg_key].add("Vehicle Data")
             
    # Cleanup corpus memory
    del full_text_corpus
    gc.collect()

    # Final List
    flows = []
    for (src, tgt, cat), data_set in connections.items():
        clean_list = sorted(list(data_set))
        flows.append({
            "from": src, "to": tgt, "category": cat, 
            "data_types": clean_list
        })
    
    flows.sort(key=lambda x: (0 if x['category'] == 'Collection' else 1, x['to']))
    return flows

def build_swimlane_diagram(entities: Dict[str, str], flows: List[Dict[str, Any]]) -> Dict[str, Any]:
    used_lanes, used_nodes = set(), set()
    node_lane_map = {}
    
    for f in flows:
        used_nodes.add(f["from"])
        used_nodes.add(f["to"])

    for node in used_nodes:
        if node == "controller": lane = "company"
        elif node == "user": lane = "user"
        else: lane = "thirdparty"
        node_lane_map[node] = lane
        used_lanes.add(lane)

    nodes = [{"id": n, "label": entities.get(n, n), "lane": node_lane_map[n]} for n in used_nodes]
    edges = [{"id": f"e{i}", "from": f["from"], "to": f["to"], "text": f"{len(f['data_types'])} Types"} for i, f in enumerate(flows)]
        
    return {"lanes": sorted(list(used_lanes)), "nodes": nodes, "edges": edges}

# ---------------------------------------------------------
# EXPORT
# ---------------------------------------------------------
def generate_mapper_report_docx(controller: str, flows: List[Dict[str, Any]], image_base64: str = None) -> bytes:
    if Document is None: raise ImportError("python-docx missing")

    doc = Document()
    doc.add_heading(f"Privacy Data Map: {controller}", level=0)
    
    if image_base64:
        try:
            img_data = base64.b64decode(image_base64.split(",")[1] if "," in image_base64 else image_base64)
            doc.add_picture(io.BytesIO(img_data), width=Inches(6.0))
        except: pass

    doc.add_heading("Data Flows", level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Action'
    hdr[1].text = 'Recipient'
    hdr[2].text = 'Data Elements'

    for flow in flows:
        row = table.add_row().cells
        row[0].text = flow['category']
        recipient_display = flow['to']
        if recipient_display == "controller": recipient_display = "Company (Internal)"
        row[1].text = recipient_display
        row[2].text = ", ".join(flow['data_types'])

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()