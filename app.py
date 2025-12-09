#!/usr/bin/env python3
import os
import re
import json
import smtplib
import io
import base64
import datetime
from typing import List, Dict, Any, Optional, Tuple
from email.message import EmailMessage

import httpx
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse
from pydantic import BaseModel

# Requires: pip install python-docx sentence-transformers numpy
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Warning: python-docx not installed. Contract features will fail.")
    Document = None

# -----------------------------------------------------------------------------
# UPGRADE: Integrated Redlining Modules (Matches your uploaded files)
# -----------------------------------------------------------------------------
try:
    # These imports map to the files: redline_apply.py, semantic_matcher.py, persona_delta.py
    from redline_apply import apply_redlines_to_docx
    from semantic_matcher import extract_paragraphs, pairwise_match
    from persona_delta import generate_document_deltas
except ImportError as e:
    print(f"Warning: Redlining modules missing or dependencies not installed: {e}")
    # Mock fallbacks so the app still boots if files are missing
    def apply_redlines_to_docx(*args, **kwargs): raise NotImplementedError("Module missing")
    def extract_paragraphs(*args): return []
    def pairwise_match(*args): return []
    async def generate_document_deltas(*args, **kwargs): return []

# ---------- Ollama + Model Config ----------

# Base URL for your Ollama instance
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")

# Default model name – you can override with:
#   export PHOENIX_MODEL_NAME=linuxlawyer
DEFAULT_MODEL_NAME = os.getenv("PHOENIX_MODEL_NAME", "llama3.2:3b")

# ---------- Jurisdiction "Personas" (Phoenix) ----------

PERSONAS: Dict[str, Dict[str, str]] = {
    "mi": {
        "label": "Michigan Laws",
        "model": DEFAULT_MODEL_NAME,
        "system": (
            "You are a legal research assistant specializing in Michigan statutory law. "
            "Your goal is to find the CLOSEST MATCHING laws to the user's query from the provided context.\n\n"
            "GUIDELINES:\n"
            "1. **Exhaustive Retrieval**: Cite EVERY relevant statute section found in the context. Do not leave out related subsections.\n"
            "2. **Strict Relevance**: Do not invent laws. If the provided context does not contain the specific answer, state that clearly.\n"
            "3. **Quote Key Language**: When you find the matching law, quote the operative text verbatim.\n"
            "4. **Jurisdiction Lock**: Focus ONLY on chunks marked [MI] (Michigan).\n\n"
            "Structure your answer with clear headings for each statute found."
        ),
    },
    "ca": {
        "label": "California Laws",
        "model": DEFAULT_MODEL_NAME,
        "system": (
            "You are a legal research assistant specializing in California statutory law. "
            "Your goal is to find the CLOSEST MATCHING laws to the user's query from the provided context.\n\n"
            "GUIDELINES:\n"
            "1. **Exhaustive Retrieval**: Cite EVERY relevant statute section found in the context. Do not leave out related subsections.\n"
            "2. **Strict Relevance**: Do not invent laws. If the provided context does not contain the specific answer, state that clearly.\n"
            "3. **Quote Key Language**: When you find the matching law, quote the operative text verbatim.\n"
            "4. **Jurisdiction Lock**: Focus ONLY on chunks marked [CA] (California).\n\n"
            "Structure your answer with clear headings for each statute found."
        ),
    },
}

# ---------- UPGRADE: Contract Review Personas (Mutable) ----------

CONTRACT_PERSONAS: Dict[str, str] = {
    "General Counsel": (
        "Tone: Professional, direct, and protective but reasonable.\n"
        "Strategy: Accept standard commercial terms. "
        "Reject unlimited liability, vague indemnities, and unilateral termination rights. "
        "Flag missing data privacy protections (GDPR/CCPA)."
    ),
    "Aggressive Litigator": (
        "Tone: Strict and demanding.\n"
        "Strategy: Maximize indemnities from the counterparty. "
        "Demand uncapped liability for IP and Confidentiality breaches. "
        "Reject any automatic renewals. Require 60-day payment terms."
    ),
    "Deal Maker": (
        "Tone: Collaborative and soft.\n"
        "Strategy: Prioritize closing the deal. "
        "Only flag critical 'bet the company' risks. "
        "Accept mutual standard clauses readily."
    )
}

# Root folder where your markdown corpus lives.
CORPUS_ROOT = os.path.expanduser("~/legal-rag")

# ---------- Optional RAG Setup (Chroma, Phoenix only) ----------

USE_RAG_BACKEND = True
RAG_DB_PATH = os.path.expanduser("~/legal-rag/db")
RAG_COLLECTION_NAME = "legal_corpus"

rag_collection = None

if USE_RAG_BACKEND:
    try:
        import chromadb
        from chromadb.utils import embedding_functions

        client = chromadb.PersistentClient(path=RAG_DB_PATH)
        embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        rag_collection = client.get_or_create_collection(
            name=RAG_COLLECTION_NAME,
            embedding_function=embedding_fn,
        )
        print(f"[RAG] Loaded Chroma collection '{RAG_COLLECTION_NAME}' from {RAG_DB_PATH}")
    except Exception as e:
        print(f"[RAG] Failed to initialize Chroma: {e}")
        rag_collection = None

# ---------- Email Config (Google Workspace SMTP Relay) ----------

SMTP_HOST = os.getenv("SMTP_HOST", "smtp-relay.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USERNAME = ""  # Not used for relay
SMTP_PASSWORD = ""  # Not used for relay
SMTP_USE_TLS = True

INTAKE_EMAIL_FROM = os.getenv("INTAKE_EMAIL_FROM", "your-email@example.com")
INTAKE_EMAIL_CC = os.getenv("INTAKE_EMAIL_CC", "your-email@example.com")

# ---------- FastAPI App ----------

app = FastAPI(title="Phoenix: Laws & Intake Engine")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- Pydantic Models: Phoenix Statute Query ----------

class QueryRequest(BaseModel):
    question: str
    personas: List[str]        # ["mi"], ["ca"], or ["mi","ca"]
    use_rag: bool = True


class PersonaAnswer(BaseModel):
    persona: str
    label: str
    answer: str


class QueryResponse(BaseModel):
    answers: List[PersonaAnswer]
    used_rag: bool
    sources: List[Dict[str, Any]]

# ---------- Pydantic Models: Intake Engine ----------

class CsuiteHit(BaseModel):
    name: str
    matched_variants: List[str] = []
    contexts: List[str] = []


class IntakeRequest(BaseModel):
    email_text: str
    reference_notes: Optional[str] = ""
    csuite_names: List[str] = []
    organization_name: Optional[str] = ""
    max_categories: int = 5  # Defaults to 5 if UI doesn't send it

    # Where to send the result (To)
    notify_email: Optional[str] = None

    # New: optional team profile config sent from the UI
    team_profile: Optional[Dict[str, Any]] = None


class IntakeResponse(BaseModel):
    categories: List[str]
    priority_label: str
    priority_score: Optional[float] = None
    summary: str
    csuite_mentions: List[CsuiteHit]

    suggested_owner: Optional[str] = None
    suggested_backup: Optional[str] = None
    suggested_next_steps: Optional[str] = None
    learning_opportunities: Optional[List[str]] = []
    # NEW — required so backend can populate these
    assigned_owner: Optional[str] = None
    assigned_backup: Optional[str] = None
    
    # NEW: Include raw input text in response
    original_text: Optional[str] = None

    raw_model_output: str
    email_status: Optional[str] = None

# ---------- Pydantic Models: Contract Redline Engine ----------

class ContractRedlineRequest(BaseModel):
    counterparty_text: str
    template_text: Optional[str] = None
    mode: str = "template_only"   # "template_only" or "hybrid_review"
    persona: str = "General Counsel"      # Name of the persona to use
    role: str = "Buyer"           # "Buyer" or "Seller"

class ContractRedlineExportRequest(BaseModel):
    original_docx_base64: str
    diff: Any 
    
class ContractReportRequest(BaseModel):
    diff: Any

class PersonaUpdateRequest(BaseModel):
    name: str
    instructions: str

# ---------- Default Team Profile (can be overridden by UI) ----------

DEFAULT_TEAM_PROFILE: Dict[str, Any] = {
    "members": [
        {
            "name": "Shawn",
            "skills": [
                {"label": "saas", "mastery": 95},
                {"label": "cybersecurity", "mastery": 95},
                {"label": "privacy", "mastery": 95},
                {"label": "contracts", "mastery": 100},
                {"label": "ai", "mastery": 100},
                {"label": "real estate", "mastery": 55},                
            ],
        },
        {
            "name": "Ron",
            "skills": [
                {"label": "negotiating", "mastery": 95},
                {"label": "contracts", "mastery": 95},
                {"label": "litigation", "mastery": 95},
                {"label": "real estate", "mastery": 95},
            ],
        },
        {
            "name": "Russell",
            "skills": [
                {"label": "open source", "mastery": 95},
                {"label": "compliance", "mastery": 95},
                {"label": "litigation", "mastery": 85},                
            ],
        },
    ]
}

# ---------- RAG Helpers (Phoenix) ----------

def infer_jurisdiction(meta: Dict[str, Any]) -> str:
    """
    Infer jurisdiction from URL or source path.
    Returns 'MI', 'CA', or 'UNK'.
    """
    url = (meta.get("url") or "").lower()
    source = (meta.get("source") or "").lower()

    if "legislature.mi.gov" in url:
        return "MI"
    if "leginfo.legislature.ca.gov" in url:
        return "CA"

    # Fallback heuristics if URL is missing:
    if "michigan" in source or "mcl" in source:
        return "MI"
    if "california" in source or "evid_" in source or "civ_" in source or "bpc_" in source:
        return "CA"

    return "UNK"


URL_RE = re.compile(r"(https?://[^\s)]+)")

# cache: source path -> {"title": ..., "url": ...}
_STATUTE_CACHE: Dict[str, Dict[str, str]] = {}


def extract_url_from_doc(doc: str, meta: Dict[str, Any]) -> str:
    """
    Fallback URL extraction from metadata or the chunk text.
    """
    def clean(u: str) -> str:
        u = u.strip()
        u = u.rstrip(").,'\"*")
        return u

    url = (meta.get("url") or "").strip()
    if url.startswith("http://") or url.startswith("https://"):
        return clean(url)

    m = URL_RE.search(doc)
    if m:
        return clean(m.group(1))

    return ""


def get_statute_info(source: str, doc: str, meta: Dict[str, Any]) -> Tuple[str, str]:
    """
    For a given 'source' (relative path stored in Chroma), read the markdown file
    once to recover a human-friendly title and a canonical statute URL.
    """
    if source in _STATUTE_CACHE:
        info = _STATUTE_CACHE[source]
        return info["title"], info["url"]

    title = None
    url = None

    def clean(u: str) -> str:
        u = u.strip()
        u = u.rstrip(").,'\"*")
        return u

    try:
        path = os.path.join(CORPUS_ROOT, source)
        with open(path, "r", encoding="utf-8") as f:
            head = f.read(4000)

        m = re.search(r"^#\s+(.*)", head, flags=re.MULTILINE)
        if m:
            title = m.group(1).strip()

        m = re.search(r"\*Statute URL:\s*(https?://\S+)", head)
        if m:
            url = clean(m.group(1))

        if not url:
            m = re.search(r"\*Group URL:\s*(https?://\S+)", head)
            if m:
                url = clean(m.group(1))

        if not url:
            m = re.search(r"\*Downloaded From:\s*(https?://\S+)", head)
            if m:
                url = clean(m.group(1))

    except Exception as e:
        print(f"[RAG] Could not read statute file for source '{source}': {e}")

    if not title:
        # Improved Fallback: Format filename to look like a title
        # e.g., "mcl_750.540.md" -> "MCL 750.540"
        base = os.path.basename(source).rsplit(".", 1)[0]
        title = base.replace("_", " ").replace("-", " ").title()

    if not url:
        url = extract_url_from_doc(doc, meta)

    if not (url.startswith("http://") or url.startswith("https://")):
        url = ""

    info = {"title": title, "url": url}
    _STATUTE_CACHE[source] = info
    return title, url

def query_chroma(question: str, n_results: int) -> Tuple[List[str], List[Dict[str, Any]]]:
    """
    Raw query to Chroma, returns (docs, metas).
    """
    result = rag_collection.query(
        query_texts=[question],
        n_results=n_results,
    )
    docs = result.get("documents", [[]])[0]
    metas = result.get("metadatas", [[]])[0]
    return docs, metas


def get_rag_context_for_persona(
    question: str,
    persona_id: str,
    k: int = 5
) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Query Chroma for relevant chunks for a SINGLE persona (jurisdiction).
    IMPROVED: Aggressively fetches and filters to find the best matches.
    Hides internal file paths from the LLM context.
    """
    if rag_collection is None:
        return "", []

    target_jur = "MI" if persona_id == "mi" else "CA" if persona_id == "ca" else None

    # Fetch a wider net initially (max 60 chunks) to ensure we capture relevant laws
    try:
        docs, metas = query_chroma(question, n_results=60)
    except Exception as e:
        print(f"[RAG] Error during query: {e}")
        return "", []

    context_pieces: List[str] = []
    sources: List[Dict[str, Any]] = []
    seen_content = set()

    idx_counter = 1
    for doc, meta in zip(docs, metas):
        # Deduplication
        if doc in seen_content:
            continue
        seen_content.add(doc)

        jur = infer_jurisdiction(meta)
        if target_jur and jur != target_jur:
            continue

        source = meta.get("source", "unknown")
        offset = meta.get("offset", 0)
        title, url = get_statute_info(source, doc, meta)

        # UPDATE: Removed 'Source: {source}' (file path) from LLM context
        # Replaced with 'Citation: {title}' to look professional
        context_pieces.append(
            f"[{idx_counter}] [{jur}] {title}\nCitation: {title}\n{doc}\n"
        )
        sources.append(
            {
                "index": idx_counter,
                "source": source, # Kept for backend ref, not shown to LLM
                "title": title,
                "url": url,
                "jurisdiction": jur,
                "offset": offset,
                "snippet": doc[:400],
            }
        )
        idx_counter += 1

        if len(context_pieces) >= k * 5:
            break

    context_text = "\n\n".join(context_pieces)
    return context_text, sources


# ---------- Contract Redline Helper (DOCX) ----------

def extract_docx_text(file_bytes: bytes) -> str:
    """
    Convert a .docx file's contents into raw text.
    """
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx module not installed on server.")
        
    try:
        buffer = io.BytesIO(file_bytes)
        doc = Document(buffer)
        lines = []
        for p in doc.paragraphs:
            lines.append(p.text)
        return "\n".join(lines).strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {e}")

def generate_analysis_report_docx(diff: List[Dict[str, Any]]) -> bytes:
    """
    Generates a clean "Findings Report" DOCX from the analysis diff.
    Structure:
      - Title
      - Clause (Original)
      - Comments (Strategy)
      - Recommended Actions (Ins/Del/Repl)
    """
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx module missing")

    doc = Document()
    
    # Title
    title = doc.add_heading('Phoenix Contract Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    doc.add_paragraph(f"Generated on: {date_str}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("---------------------------------------------------------")

    if not diff:
        doc.add_paragraph("No issues or redlines found in this analysis.")
        
    for i, item in enumerate(diff, 1):
        cp_text = item.get("cp_text", "")
        delta = item.get("delta", {})
        
        # Extract fields
        comments = delta.get("comments", [])
        insertions = delta.get("insertions", [])
        deletions = delta.get("deletions", [])
        replacements = delta.get("replacements", [])
        
        # Skip if empty (double check)
        if not (comments or insertions or deletions or replacements):
            continue

        # Header for the clause
        snippet = (cp_text[:60] + "...") if len(cp_text) > 60 else cp_text
        h = doc.add_heading(f"Clause {i}: {snippet}", level=2)
        
        # Original Text Block
        p_orig = doc.add_paragraph()
        run_orig = p_orig.add_run(cp_text)
        run_orig.font.italic = True
        run_orig.font.color.rgb = RGBColor(100, 100, 100) # Grey

        # Comments Section
        if comments:
            doc.add_heading("Strategy / Comments:", level=3)
            for c in comments:
                doc.add_paragraph(c, style='List Bullet')

        # Changes Section
        if insertions or deletions or replacements:
            doc.add_heading("Recommended Changes:", level=3)
            
            for ins in insertions:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"INSERT: {ins}")
                run.font.color.rgb = RGBColor(0, 128, 0) # Green
                
            for dele in deletions:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"DELETE: {dele}")
                run.font.color.rgb = RGBColor(200, 0, 0) # Red

            for rep in replacements:
                if isinstance(rep, dict):
                    frm = rep.get("from", "")
                    to = rep.get("to", "")
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run("CHANGE: ")
                    p.add_run(f'"{frm}"').font.strike = True
                    p.add_run("  -->  ")
                    p.add_run(f'"{to}"').font.bold = True

        doc.add_paragraph("_" * 50) # Divider

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- Ollama Helper (JSON MODE) ----------

async def call_ollama_generate(model: str, prompt: str, json_mode: bool = False) -> str:
    url = f"{OLLAMA_URL}/api/generate"
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "num_predict": 1024,
    }
    if json_mode:
        payload["format"] = "json"

    # CHANGE: Increased timeout from 180 to 600 (10 minutes) for large contracts
    async with httpx.AsyncClient(timeout=600.0) as client:
        resp = await client.post(url, json=payload)
        resp.raise_for_status()
        return resp.json().get("response", "").strip()


def build_prompt(persona_id: str, question: str, context: Optional[str]) -> str:
    """
    Compose the full prompt for a given jurisdiction persona, question, and optional RAG context.
    """
    persona = PERSONAS[persona_id]
    system = persona["system"]

    base = system + "\n\n"

    if context:
        base += (
            "You have the following statutory materials found in the database. "
            "These are your ONLY source of truth for specific statutes.\n"
            "----- BEGIN CONTEXT -----\n"
            f"{context}\n"
            "----- END CONTEXT -----\n\n"
        )

    base += (
        "TASK:\n"
        "1. Answer the user's question using the provided context.\n"
        "2. If the context contains the specific statute requested, QUOTE IT DIRECTLY.\n"
        "3. If the context contains relevant but not exact matches, list them as 'Related Laws'.\n"
        "4. If no relevant laws are found in the context, explicitly say 'I could not find a specific statute in the database matching this query.'\n\n"
        f"User question:\n{question}\n\n"
        "Answer:\n"
    )
    return base


# ---------- Intake Prompt Builder ----------
def build_intake_prompt(req: IntakeRequest) -> str:
    """
    Build a prompt for the intake engine to:
      - classify category
      - assign priority
      - detect C-suite names
      - summarize & recommend next steps
    """
    # --- UPDATED PRIORITY RUBRIC ---
    priority_instructions = (
        "PRIORITY RUBRIC:\n"
        "- **Critical (10)**: Immediate financial risk, data breach, lawsuit served, or 'Urgent' from CEO.\n"
        "- **High (8)**: Deadlines within 24-48h, regulatory compliance, important deal closings.\n"
        "- **Medium (5)**: Standard contract reviews (NDAs, MSAs), routine advisory questions.\n"
        "- **Low (2)**: Administrative tasks, scheduling, FYI notifications, no deadline.\n"
    )

    base = (
        "You are an intake triage assistant. You do NOT give legal advice.\n"
        "You classify internal matters, assign priority, flag C-suite mentions, "
        "summarize, and recommend owners based on the notes provided.\n\n"
        f"{priority_instructions}\n\n"
        "Respond ONLY with strict JSON matching this schema:\n"
        "{\n"
        '  "categories": ["..."],\n'
        '  "priority_label": "Critical | High | Medium | Low",\n'
        '  "priority_score": number,\n'
        '  "summary": "string",\n'
        '  "csuite_mentions": [\n'
        '    {"name": "string", "matched_variants": ["..."], "contexts": ["..."]}\n'
        "  ],\n"
        '  "suggested_owner": "string",\n'
        '  "suggested_next_steps": "string"\n'
        "}\n\n"
    )

    if req.organization_name:
        base += f"Organization: {req.organization_name}\n\n"

    if req.csuite_names:
        joined = ", ".join(req.csuite_names)
        base += (
            f"C-suite names to detect: {joined}.\n"
            "Treat role-references (e.g., 'the CEO') as matches.\n\n"
        )

    if req.reference_notes:
        base += (
            "Reference Notes / Playbook (Use this to determine suggested_owner if specified):\n"
            "----- BEGIN NOTES -----\n"
            f"{req.reference_notes}\n"
            "----- END NOTES -----\n\n"
        )

    base += (
        "Email to classify:\n"
        "----- BEGIN EMAIL -----\n"
        f"{req.email_text}\n"
        "----- END EMAIL -----\n\n"
        f"Limit categories to max {req.max_categories}.\n"
        "Respond now in strict JSON only.\n"
    )

    return base

def _safe_parse_intake_json(model_output: str) -> Dict[str, Any]:
    """
    Extract JSON object from the LLM output and parse it safely.
    Returns fallback values if parsing fails.
    """
    # 1. Try to parse directly (if JSON mode worked perfectly)
    try:
        return json.loads(model_output)
    except:
        pass

    # 2. Fallback: Extract via regex
    try:
        m = re.search(r"\{.*\}", model_output, re.DOTALL)
        if m:
            return json.loads(m.group(0))
    except Exception as e:
        print(f"[Intake] Failed to parse JSON: {e}")

    return {
        "categories": [],
        "priority_label": "Unknown",
        "priority_score": None,
        "summary": "",
        "csuite_mentions": [],
        "suggested_owner": None,
        "suggested_next_steps": None,
    }

def _skill_match_score(skill_label: str, category_label: str) -> float:
    """
    Very small fuzzy matcher between a skill label and a category label.
    """
    s = skill_label.lower().strip()
    c = category_label.lower().strip()
    if not s or not c:
        return 0.0

    tokens_s = set(s.split())
    tokens_c = set(c.split())

    # Strong match: any shared word
    if tokens_s & tokens_c:
        return 1.0

    # Substring match
    if s in c or c in s:
        return 0.8

    # Jaccard similarity as a softer signal
    union = tokens_s | tokens_c
    if not union:
        return 0.0
    jaccard = len(tokens_s & tokens_c) / len(union)
    if jaccard >= 0.30:
        return 0.6

    return 0.0


def assign_team_owner(
    result: IntakeResponse,
    team_profile: Optional[Dict[str, Any]]
) -> IntakeResponse:
    """
    Assign lead + backup.
    """
    profile = team_profile or DEFAULT_TEAM_PROFILE
    members = profile.get("members") or []

    if not members:
        result.assigned_owner = None
        result.assigned_backup = None
        result.learning_opportunities = []
        result.suggested_owner = None
        result.suggested_backup = None
        return result

    member_map = {m.get("name", "").strip().lower(): m.get("name", "").strip() for m in members}
    categories = [c.lower() for c in (result.categories or [])]

    # --- 1. PLAYBOOK OVERRIDE CHECK ---
    llm_suggestion = (result.suggested_owner or "").strip().lower()
    llm_suggestion_clean = llm_suggestion.split("(")[0].strip()
    
    forced_lead_name = None

    if llm_suggestion_clean in member_map:
        forced_lead_name = member_map[llm_suggestion_clean]
        result.assigned_owner = forced_lead_name
        result.assigned_backup = None
        result.suggested_owner = f"{forced_lead_name} (Playbook Override)"
        result.suggested_backup = "N/A"
        # We continue to calculate training opportunities below based on the forced lead

    # --- 2. SKILL MATH ROUTING (if no override) ---
    lead_candidates: List[Tuple[float, str]] = []   # (score, name)
    training_set: set[str] = set()

    for member in members:
        name = (member.get("name") or "").strip()
        if not name:
            continue

        skills = member.get("skills") or []
        lead_score = 0.0
        has_skill_match = False

        for skill in skills:
            label = (skill.get("label") or "").strip()
            if not label:
                continue

            mastery = float(skill.get("mastery") or 0)
            
            # Compute best match across all categories for this skill
            best_match = 0.0
            for cat in categories:
                best_match = max(best_match, _skill_match_score(label, cat))

            if best_match > 0.0:
                has_skill_match = True
                # Score calculation
                mastery_weight = 0.5 + (max(0.0, min(mastery, 100.0)) / 200.0)
                lead_score += best_match * mastery_weight

        if lead_score > 0.0:
            lead_candidates.append((lead_score, name))
        
        # Determine if they are a training candidate (will filter Lead out later)
        if has_skill_match:
            training_set.add(name)

    # Assign Lead/Backup if not forced
    if not forced_lead_name:
        if not lead_candidates:
            result.assigned_owner = members[0].get("name")
            result.assigned_backup = None
        else:
            lead_candidates.sort(key=lambda x: x[0], reverse=True)
            result.assigned_owner = lead_candidates[0][1]
            result.assigned_backup = lead_candidates[1][1] if len(lead_candidates) > 1 else None

        result.suggested_owner = f"{result.assigned_owner} (lead)" if result.assigned_owner else "N/A"
        result.suggested_backup = (
            f"{result.assigned_backup} (backup)" if result.assigned_backup else "N/A"
        )

    # --- 3. TRAINING ASSIGNMENT ---
    # Training = Anyone with a skill match who is NOT the assigned lead
    final_lead = result.assigned_owner
    
    final_training_list = []
    for candidate in training_set:
        if candidate != final_lead:
            final_training_list.append(candidate)
            
    result.learning_opportunities = sorted(final_training_list)

    return result


def _build_intake_response(
    parsed: Dict[str, Any], 
    raw: str, 
    original_text: str,
    watchlist: List[str] = None
) -> IntakeResponse:
    """
    Convert parsed JSON into IntakeResponse Pydantic model.
    Includes Dynamic Priority Mapping and Strict C-Suite Filtering.
    """

    # Extract categories safely
    categories = parsed.get("categories") or []
    if not isinstance(categories, list):
        categories = [str(categories)]

    # --- DYNAMIC PRIORITY MAPPING (Strict Sync) ---
    priority_label_raw = str(parsed.get("priority_label") or "").strip().lower()
    
    # Defaults
    priority_label = "Medium"
    priority_score = 5.0
    
    # 1. Map known synonyms for Label
    if "critical" in priority_label_raw or "urgent" in priority_label_raw:
        priority_label = "Critical"
        priority_score = 10.0
    elif "high" in priority_label_raw:
        priority_label = "High"
        priority_score = 8.0
    elif "medium" in priority_label_raw or "normal" in priority_label_raw or "standard" in priority_label_raw:
        priority_label = "Medium"
        priority_score = 5.0
    elif "low" in priority_label_raw or "routine" in priority_label_raw:
        priority_label = "Low"
        priority_score = 2.0
    else:
        # 2. Fallback: Check if score exists and reverse-map it
        try:
            score_val = float(parsed.get("priority_score") or 0)
        except:
            score_val = 0
            
        if score_val >= 9:
            priority_label = "Critical"
            priority_score = 10.0
        elif score_val >= 7:
            priority_label = "High"
            priority_score = 8.0
        elif score_val >= 4:
            priority_label = "Medium"
            priority_score = 5.0
        else:
            priority_label = "Low"
            priority_score = 2.0

    summary = parsed.get("summary") or ""

    # --- C-SUITE FILTERING LOGIC ---
    csuite_mentions_raw = parsed.get("csuite_mentions") or []
    csuite_mentions = []
    
    # Normalize Watchlist (if provided)
    normalized_watchlist = [w.lower().strip() for w in (watchlist or []) if w.strip()]
    
    # FIX: If no watchlist is provided, we ignore all LLM detections to prevent hallucinations.
    if normalized_watchlist:
        if isinstance(csuite_mentions_raw, list):
            for item in csuite_mentions_raw:
                if not isinstance(item, dict):
                    continue
                    
                name = item.get("name") or ""
                det_lower = name.lower().strip()
                
                # Check match against watchlist
                is_valid = False
                for w in normalized_watchlist:
                    # Match if watchlist item appears in detection (e.g. "CEO" in "The CEO")
                    # OR if detection matches watchlist (e.g. "Ron" in "Ron Conway")
                    if w in det_lower or det_lower in w:
                        is_valid = True
                        break
                
                if not is_valid:
                    continue # Skip this hit
                
                matched_variants = item.get("matched_variants") or []
                contexts = item.get("contexts") or []
                if not isinstance(matched_variants, list):
                    matched_variants = [str(matched_variants)]
                if not isinstance(contexts, list):
                    contexts = [str(contexts)]
                    
                csuite_mentions.append(
                    CsuiteHit(
                        name=name,
                        matched_variants=[str(v) for v in matched_variants],
                        contexts=[str(c) for c in contexts],
                    )
                )
    else:
        # No watchlist provided = no detection allowed
        csuite_mentions = []

    # --- CATEGORY CLEANUP ---
    # If no C-suite mentions survived filtering, ensure "C-Suite" isn't in categories
    if not csuite_mentions:
        # Remove tags that imply executive presence if we have no hits
        sensitive_tags = ["c-suite", "executive", "executive mention", "urgent - executive"]
        cleaned_categories = []
        for cat in categories:
            cat_lower = str(cat).lower().strip()
            # If the category tag contains "c-suite" or "executive", drop it
            if any(bad in cat_lower for bad in sensitive_tags):
                continue
            cleaned_categories.append(cat)
        categories = cleaned_categories

    suggested_owner = parsed.get("suggested_owner")
    suggested_next_steps = parsed.get("suggested_next_steps")

    return IntakeResponse(
        categories=[str(c) for c in categories],
        priority_label=priority_label.title(),
        priority_score=float(priority_score),
        summary=str(summary),
        csuite_mentions=csuite_mentions,
        suggested_owner=str(suggested_owner) if suggested_owner else None,
        suggested_next_steps=str(suggested_next_steps) if suggested_next_steps else None,
        raw_model_output=raw,
        email_status=None,
        original_text=original_text
    )

# ---------- EMAIL SENDING (Google Workspace SMTP Relay) ----------

def send_intake_email(
    to_email: str,
    result: IntakeResponse,
    original_email_text: Optional[str] = None,
) -> str:
    """
    Sends intake analysis results.
    """
    if not SMTP_HOST:
        return "email_not_configured"

    try:
        categories = ", ".join(result.categories) if result.categories else "Uncategorized"
        subject = f"[Phoenix Intake] {result.priority_label} – {categories}"

        lines: List[str] = []
        lines.append(f"Priority: {result.priority_label}")

        if isinstance(result.priority_score, (int, float)):
            lines.append(f"Score: {result.priority_score:.1f}/10")
        else:
            lines.append("Score: N/A")

        lines.append(f"Categories: {categories}")
        lines.append("")

        lines.append("Summary:")
        lines.append(result.summary or "(none)")
        lines.append("")

        lines.append(f"Suggested owner: {result.suggested_owner or 'N/A'}")
        lines.append(f"Suggested backup: {result.suggested_backup or 'N/A'}")

        if getattr(result, "learning_opportunities", None):
            learner_line = ", ".join(result.learning_opportunities)
            lines.append(f"Training Opportunity: {learner_line}")
        else:
            lines.append("Training Opportunity: N/A")

        lines.append("")
        lines.append("Suggested next steps:")
        lines.append(result.suggested_next_steps or "(none)")

        if result.csuite_mentions:
            lines.append("")
            lines.append("C-suite mentions:")
            for hit in result.csuite_mentions:
                variants = ", ".join(hit.matched_variants) if hit.matched_variants else "-"
                lines.append(f" - {hit.name} (variants: {variants})")

        if original_email_text:
            lines.append("")
            lines.append("Inbound request (raw text):")
            lines.append("----------------------------------------")
            lines.append(original_email_text)
            lines.append("----------------------------------------")

        lines.append("")
        lines.append("Raw model output:")
        lines.append(result.raw_model_output)

        msg = EmailMessage()
        msg["Subject"] = subject
        msg["From"] = INTAKE_EMAIL_FROM
        msg["To"] = to_email
        msg["Cc"] = INTAKE_EMAIL_CC
        msg.set_content("\n".join(lines))

        recipients = [to_email, INTAKE_EMAIL_CC]

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            if SMTP_PORT == 587:
                server.starttls()
            if SMTP_USERNAME and SMTP_PASSWORD:
                server.login(SMTP_USERNAME, SMTP_PASSWORD)
            server.send_message(msg, to_addrs=recipients)

        return "sent"

    except Exception as e:
        print(f"Email send error: {e}")
        return f"error: {e}"

# ---------- FASTAPI ROUTES ----------

@app.get("/api/health")
async def health():
    return {"status": "ok"}

# ---------- API Routes: Personas (NEW) ----------

@app.get("/api/contracts/personas")
async def get_contract_personas():
    """Return list of available contract personas."""
    return [{"name": k, "instructions": v} for k, v in CONTRACT_PERSONAS.items()]

@app.post("/api/contracts/personas")
async def upsert_contract_persona(req: PersonaUpdateRequest):
    """Add or update a persona."""
    CONTRACT_PERSONAS[req.name] = req.instructions
    return {"status": "ok", "personas": await get_contract_personas()}

@app.delete("/api/contracts/personas/{name}")
async def delete_contract_persona(name: str):
    if name in CONTRACT_PERSONAS:
        del CONTRACT_PERSONAS[name]
    return {"status": "ok", "personas": await get_contract_personas()}


# ---------- API Routes: Redline ----------

@app.post("/api/contracts/redline/upload")
async def upload_contracts_for_redline(
    counterparty: UploadFile = File(...),
    template: UploadFile = File(None),
):
    try:
        if not counterparty.filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="Counterparty file must be a .docx")
        
        cp_bytes = await counterparty.read()
        cp_text = extract_docx_text(cp_bytes)
        
        tp_text = None
        if template:
            tp_bytes = await template.read()
            tp_text = extract_docx_text(tp_bytes)

        return {"status": "ok", "counterparty_text": cp_text, "template_text": tp_text}

    except Exception as e:
        print(f"[Upload] Failed: {e}")
        raise HTTPException(status_code=500, detail=f"Upload Failed: {str(e)}")

@app.post("/api/contracts/redline/analyze")
async def analyze_contract_redline(req: ContractRedlineRequest):
    try:
        # Debug Log
        print(f"[Analyze] Starting analysis. Persona='{req.persona}', Role='{req.role}'")
        
        # 1. Clean & Extract Paragraphs
        cp_paragraphs = extract_paragraphs(req.counterparty_text)
        tp_paragraphs = extract_paragraphs(req.template_text) if req.template_text else []
        
        # 2. Semantic Matching
        matched_pairs = pairwise_match(cp_paragraphs, tp_paragraphs, threshold=0.45)
        print(f"[Analyze] Matches found: {len(matched_pairs)}")

        # 3. Retrieve Persona Instructions (Dynamic Lookup)
        # Use provided persona name, fallback to a default string if missing/deleted
        persona_instructions = CONTRACT_PERSONAS.get(
            req.persona, 
            CONTRACT_PERSONAS.get("General Counsel", "Act as a helpful lawyer.")
        )

        # 4. LLM Wrapper
        async def llm_wrapper(prompt, json_mode=False):
            return await call_ollama_generate(DEFAULT_MODEL_NAME, prompt, json_mode=json_mode)

        # 5. Generate Deltas (Passing Role)
        print("[Analyze] Calling LLM for redlines...")
        
        # Updated to use Turbo Mode (parallel) implicitly via concurrency limit
        deltas = await generate_document_deltas(
            llm_generate_fn=llm_wrapper,
            matched_paragraphs=matched_pairs,
            persona_instructions=persona_instructions,
            role=req.role  # Pass role to the engine
        )
        print("[Analyze] LLM generation complete.")

        # --- NEW: Filter out empty/ignored clauses ---
        final_diff = []
        for d in deltas:
            delta_content = d.get("delta", {})
            has_changes = (
                delta_content.get("insertions") or 
                delta_content.get("deletions") or 
                delta_content.get("replacements") or 
                delta_content.get("comments")
            )
            if has_changes:
                final_diff.append(d)

        # --- FIX: Use 'matched_pairs' instead of 'matched_paragraphs' ---
        return {"status": "ok", "mode": req.mode, "diff": final_diff, "match_count": len(matched_pairs)}

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Server Error: {str(e)}")

@app.post("/api/contracts/redline/export")
async def export_contract_redline(req: ContractRedlineExportRequest):
    """
    Step 3: Takes original file (base64) and the LLM diff, applies changes via Python-Docx / XML editing,
    and returns a downloadable tracked-changes DOCX file.
    """
    original_bytes = base64.b64decode(req.original_docx_base64)
    diff = req.diff

    try:
        redlined_bytes = apply_redlines_to_docx(original_bytes, diff)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX redline: {e}")

    return StreamingResponse(
        io.BytesIO(redlined_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=redlined.docx"}
    )
    
@app.post("/api/contracts/redline/export-report")
async def export_contract_report(req: ContractReportRequest):
    """
    NEW: Generates a human-readable 'Summary Report' DOCX listing all findings.
    """
    try:
        report_bytes = generate_analysis_report_docx(req.diff)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate report: {e}")

    return StreamingResponse(
        io.BytesIO(report_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Analysis_Report.docx"}
    )


@app.post("/api/legal/query", response_model=QueryResponse)
async def legal_query(req: QueryRequest):
    """
    Main endpoint for Phoenix: multi-persona statute exploration.
    """
    requested = [p for p in req.personas if p in PERSONAS]
    if not requested:
        requested = ["mi"]

    persona_contexts = {}
    all_sources = []
    used_rag = False

    # RAG lookup (per persona)
    if req.use_rag and rag_collection is not None:
        for persona_id in requested:
            ctx, srcs = get_rag_context_for_persona(req.question, persona_id, k=5)
            if ctx:
                persona_contexts[persona_id] = ctx
            all_sources.extend(srcs)

        if all_sources:
            used_rag = True

        # Deduplicate
        seen = set()
        deduped = []
        for s in all_sources:
            key = (s["source"], s["jurisdiction"], s["title"])
            if key in seen:
                continue
            seen.add(key)
            deduped.append(s)
        all_sources = deduped

    # Generate answers
    answers = []
    for persona_id in requested:
        persona = PERSONAS[persona_id]
        ctx = persona_contexts.get(persona_id, "")
        prompt = build_prompt(persona_id, req.question, ctx if used_rag else None)

        try:
            # Note: We do NOT use JSON mode here; we want free text.
            reply = await call_ollama_generate(persona["model"], prompt, json_mode=False)
        except Exception as e:
            reply = f"(Error: {e})"

        answers.append(
            PersonaAnswer(
                persona=persona_id,
                label=persona["label"],
                answer=reply,
            )
        )

    return QueryResponse(
        answers=answers,
        used_rag=used_rag,
        sources=all_sources,
    )


@app.post("/api/intake/analyze", response_model=IntakeResponse)
async def intake_analyze(req: IntakeRequest):
    """
    Analyze an inbound email/message and optionally email the result.
    ALWAYS returns a valid IntakeResponse — never None.
    """

    prompt = build_intake_prompt(req)

    # ------------------------------------------------------------------
    # Call model — FORCE JSON MODE (Fixes parsing errors)
    # ------------------------------------------------------------------
    try:
        raw_output = await call_ollama_generate(DEFAULT_MODEL_NAME, prompt, json_mode=True)
    except Exception as e:
        return IntakeResponse(
            categories=[],
            priority_label="Error",
            priority_score=None,
            summary=f"Error contacting model: {e}",
            csuite_mentions=[],
            suggested_owner=None,
            suggested_next_steps=None,
            raw_model_output=str(e),
            email_status="skipped",
            assigned_owner=None,
            assigned_backup=None,
        )

    # ------------------------------------------------------------------
    # Parse JSON (safe)
    # ------------------------------------------------------------------
    parsed = _safe_parse_intake_json(raw_output)

    # Build typed model
    # NEW: Passing req.email_text as original_text
    # NEW: Passing req.csuite_names as watchlist for strict filtering
    result = _build_intake_response(parsed, raw_output, req.email_text, req.csuite_names)
    
    # ---- NEW: enforce strict team routing using UI config ----
    result = assign_team_owner(result, req.team_profile)


    # ------------------------------------------------------------------
    # Guarantee all required fields exist
    # ------------------------------------------------------------------
    if not hasattr(result, "assigned_owner"):
        result.assigned_owner = None
    if not hasattr(result, "assigned_backup"):
        result.assigned_backup = None

    # ------------------------------------------------------------------
    # Email sending
    # ------------------------------------------------------------------
    email_status = None
    if req.notify_email:
        try:
            # NEW: Passing the original email text to the email function
            email_status = send_intake_email(req.notify_email, result, original_email_text=req.email_text)
        except Exception as e:
            email_status = f"error: {e}"

    result.email_status = email_status

    # ------------------------------------------------------------------
    # CRITICAL: ALWAYS RETURN RESULT
    # ------------------------------------------------------------------
    return result

# ---------- Phoenix UI at /ui ----------

@app.get("/ui", response_class=HTMLResponse)
async def ui():
    html = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Phoenix Laws</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&display=swap" rel="stylesheet">
  <style>
    :root {
      --primary: #90CAF9;
      --primary-dark: #42A5F5;
      --bg: #121212;
      --surface: #1E1E1E;
      --surface-2: #2C2C2C;
      --text: #E0E0E0;
      --text-secondary: #B0B0B0;
      --border: #333;
      --success: #66BB6A;
    }
    body {
      font-family: 'Roboto', sans-serif;
      background-color: var(--bg);
      color: var(--text);
      margin: 0;
      padding: 0;
      line-height: 1.6;
    }
    /* App Bar */
    .app-bar {
      background-color: var(--surface);
      padding: 0 24px;
      height: 64px;
      display: flex;
      align-items: center;
      justify-content: space-between;
      box-shadow: 0 2px 4px rgba(0,0,0,0.3);
      position: sticky;
      top: 0;
      z-index: 100;
    }
    .app-bar-title {
      font-size: 1.25rem;
      font-weight: 500;
      color: var(--primary);
    }
    .nav-links a {
      color: var(--text-secondary);
      text-decoration: none;
      margin-left: 20px;
      font-size: 0.9rem;
      transition: color 0.2s;
    }
    .nav-links a:hover, .nav-links a.active {
      color: var(--primary);
    }

    /* Main Container */
    .container {
      max-width: 900px;
      margin: 24px auto;
      padding: 0 16px;
    }
    .card {
      background: var(--surface);
      border-radius: 8px;
      padding: 24px;
      box-shadow: 0 1px 3px rgba(0,0,0,0.2), 0 2px 8px rgba(0,0,0,0.1);
      margin-bottom: 24px;
      transition: box-shadow 0.3s ease;
    }
    .card:hover {
       box-shadow: 0 4px 6px rgba(0,0,0,0.3), 0 8px 16px rgba(0,0,0,0.1);
    }
    h2 { font-weight: 400; font-size: 1.5rem; margin-top: 0; margin-bottom: 8px; }
    p.subtitle { color: var(--text-secondary); font-size: 0.9rem; margin-bottom: 24px; margin-top: 0; }

    /* Inputs */
    .input-group { position: relative; margin-bottom: 20px; }
    textarea {
      width: 100%;
      background: rgba(255,255,255,0.05);
      border: 1px solid var(--border);
      border-radius: 4px;
      color: var(--text);
      padding: 16px;
      font-family: inherit;
      font-size: 1rem;
      min-height: 120px;
      resize: vertical;
      box-sizing: border-box;
      transition: border-color 0.2s;
    }
    textarea:focus { outline: none; border-color: var(--primary); background: rgba(255,255,255,0.08); }
    label { display: block; margin-bottom: 8px; color: var(--text-secondary); font-size: 0.85rem; font-weight: 500; }

    /* Controls Row */
    .controls { display: flex; align-items: center; gap: 24px; flex-wrap: wrap; margin-bottom: 24px; }
    
    /* Chips (Checkbox) */
    .chip-group { display: flex; gap: 12px; }
    .chip-input { display: none; }
    .chip-label {
      background: var(--surface-2);
      padding: 8px 16px;
      border-radius: 16px;
      font-size: 0.9rem;
      cursor: pointer;
      border: 1px solid transparent;
      transition: all 0.2s;
      user-select: none;
    }
    .chip-input:checked + .chip-label {
      background: rgba(144, 202, 249, 0.15);
      color: var(--primary);
      border-color: var(--primary);
    }

    /* Switch */
    .switch-label { display: flex; align-items: center; gap: 12px; cursor: pointer; font-size: 0.9rem; }
    .switch {
      position: relative; width: 36px; height: 20px; background: #555; border-radius: 20px; transition: 0.3s;
    }
    .switch::after {
      content: ''; position: absolute; top: 2px; left: 2px; width: 16px; height: 16px; 
      background: #fff; border-radius: 50%; transition: 0.3s;
    }
    input:checked + .switch-label .switch { background: var(--primary-dark); }
    input:checked + .switch-label .switch::after { transform: translateX(16px); }

    /* Button */
    .btn {
      background: var(--primary);
      color: #000;
      border: none;
      padding: 10px 24px;
      border-radius: 4px;
      font-size: 0.95rem;
      font-weight: 500;
      text-transform: uppercase;
      letter-spacing: 0.5px;
      cursor: pointer;
      box-shadow: 0 2px 4px rgba(0,0,0,0.2);
      transition: filter 0.2s, box-shadow 0.2s;
    }
    .btn:hover { filter: brightness(1.1); box-shadow: 0 4px 8px rgba(0,0,0,0.3); }
    .btn:disabled { opacity: 0.6; cursor: default; }

    /* Results */
    .answer-card {
      background: var(--surface);
      border-radius: 8px;
      padding: 0;
      overflow: hidden;
      margin-top: 24px;
      border: 1px solid var(--border);
    }
    .answer-header {
      background: rgba(255,255,255,0.03);
      padding: 12px 20px;
      border-bottom: 1px solid var(--border);
      display: flex; justify-content: space-between; align-items: center;
    }
    .badge {
      font-size: 0.75rem; padding: 2px 8px; border-radius: 12px; 
      background: rgba(102, 187, 106, 0.2); color: var(--success); border: 1px solid rgba(102, 187, 106, 0.4);
    }
    .answer-body {
      padding: 20px;
      white-space: pre-wrap;
      font-family: 'Roboto', sans-serif; /* Override monospace unless specific */
      font-size: 0.95rem;
      color: #dcdcdc;
    }
    .sources-box {
      margin-top: 16px;
      padding: 16px;
      background: #151515;
      border-top: 1px solid var(--border);
      font-size: 0.85rem;
    }
    .source-link { color: var(--primary); text-decoration: none; }
    .source-link:hover { text-decoration: underline; }
    
    .status-text { color: var(--text-secondary); font-size: 0.9rem; margin-top: 8px; font-style: italic; }
  </style>
</head>
<body>
  <div class="app-bar">
    <div class="app-bar-title">Phoenix Laws</div>
    <div class="nav-links">
      <a href="/ui" class="active">Laws</a>
      <a href="/ui/intake">Intake</a>
      <a href="/ui/contracts">Contracts</a>
    </div>
  </div>

  <div class="container">
    <div class="card">
      <h2>Legal Research (Laws)</h2>
      <p class="subtitle">Select jurisdictions to retrieve relevant statutory text and explanations.</p>
      
      <div class="input-group">
        <label for="question">YOUR QUESTION</label>
        <textarea id="question" placeholder="e.g., What are the requirements for data breach notification in Michigan?"></textarea>
      </div>

      <div class="controls">
        <div>
            <label>JURISDICTIONS</label>
            <div class="chip-group">
                <label>
                    <input type="checkbox" id="p_mi" class="chip-input" checked>
                    <span class="chip-label">Michigan</span>
                </label>
                <label>
                    <input type="checkbox" id="p_ca" class="chip-input">
                    <span class="chip-label">California</span>
                </label>
            </div>
        </div>
        
        <div style="margin-top:24px;"> <input type="checkbox" id="use_rag" style="display:none;">
           <label for="use_rag" class="switch-label">
             <div class="switch"></div>
             <span>Use Knowledge Base (RAG)</span>
           </label>
        </div>
      </div>

      <button id="ask_btn" class="btn">Pull Statutes</button>
      <div id="status" class="status-text"></div>
    </div>

    <div id="answers"></div>
    
    <div style="text-align:center; color: #555; font-size: 0.75rem; margin-top: 40px;">
        CONFIDENTIAL &bull; INTERNAL DEMO ONLY &bull; NOT LEGAL ADVICE
    </div>
  </div>

  <script>
    async function askAgents() {
      const btn = document.getElementById("ask_btn");
      const status = document.getElementById("status");
      const answersDiv = document.getElementById("answers");
      const q = document.getElementById("question").value.trim();
      const useRag = document.getElementById("use_rag").checked;

      const personas = [];
      if (document.getElementById("p_mi").checked) personas.push("mi");
      if (document.getElementById("p_ca").checked) personas.push("ca");

      if (!q) {
        status.textContent = "Please enter a question.";
        return;
      }
      if (personas.length === 0) {
        status.textContent = "Select at least one jurisdiction.";
        return;
      }

      btn.disabled = true;
      status.textContent = "Searching statutes and generating response...";
      answersDiv.innerHTML = "";

      try {
        const resp = await fetch("/api/legal/query", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({
            question: q,
            personas: personas,
            use_rag: useRag
          })
        });
        if (!resp.ok) {
          const t = await resp.text();
          status.textContent = "Error: " + t;
          btn.disabled = false;
          return;
        }
        const data = await resp.json();
        status.textContent = "";

        data.answers.forEach(ans => {
          const card = document.createElement("div");
          card.className = "answer-card";

          const header = document.createElement("div");
          header.className = "answer-header";
          
          header.innerHTML = `
            <span style="font-weight:500; color:white;">${ans.label}</span>
            <span class="badge">${data.used_rag ? "RAG ACTIVE" : "NO RAG"}</span>
          `;

          const body = document.createElement("div");
          body.className = "answer-body";
          body.textContent = ans.answer;

          card.appendChild(header);
          card.appendChild(body);
          answersDiv.appendChild(card);
        });

        if (data.used_rag && data.sources && data.sources.length) {
          const srcDiv = document.createElement("div");
          srcDiv.className = "sources-box";
          let html = "<div style='color:var(--text-secondary); margin-bottom:12px; font-weight:500;'>CITATIONS & SOURCES</div>";
          
          data.sources.forEach(s => {
             // Use Title. If missing, clean up the source path (fallback).
             let displayTitle = s.title;
             if (!displayTitle) {
                  displayTitle = s.source.split('/').pop().replace('.md', '').replace(/_/g, ' ').toUpperCase();
             }
             
             const label = s.jurisdiction === "MI" ? "Michigan" : s.jurisdiction === "CA" ? "California" : "Ref";
             
             // Construct the link (Action)
             let action = "";
             if (s.url && s.url.startsWith("http")) {
                 action = `<a href="${s.url}" class="source-link" target="_blank">[Official Source]</a>`;
             } else {
                 action = `<span style="font-size:0.8rem; color:#666;">(No online link available)</span>`;
             }
             
             // RENDER BLOCK: Clean Title + Link below it. No file path.
             html += `<div style="margin-bottom:12px; font-family:'Roboto', sans-serif; font-size:0.9rem; border-left:3px solid #444; padding-left:12px;">
                <div style="font-weight:500; color:#e0e0e0;">${displayTitle} <span style="font-size:0.75em; color:#888; margin-left:8px; text-transform:uppercase;">${label}</span></div>
                <div style="margin-top:2px;">${action}</div>
             </div>`;
          });
          srcDiv.innerHTML = html;
          answersDiv.appendChild(srcDiv);
        }

      } catch (err) {
        console.error(err);
        status.textContent = "Network error: " + err;
      } finally {
        btn.disabled = false;
      }
    }

    document.getElementById("ask_btn").addEventListener("click", askAgents);
    // Initialize switch default
    document.getElementById("use_rag").checked = true;
  </script>
</body>
</html>
    """
    return HTMLResponse(content=html)


# ---------- Intake UI at /ui/intake ----------

@app.get("/ui/intake", response_class=HTMLResponse)
async def intake_ui():
    html = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Phoenix Intake Engine</title>
  <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&family=Roboto+Mono:wght@400;500&display=swap" rel="stylesheet">
  <style>
    :root {
      --primary: #90CAF9;
      --bg: #121212;
      --surface: #1E1E1E;
      --surface-2: #2C2C2C;
      --text: #E0E0E0;
      --text-sec: #A0A0A0;
      --border: #333;
      --accent: #BB86FC;
    }
    * { box-sizing: border-box; }
    body {
      font-family: 'Roboto', sans-serif;
      background-color: var(--bg);
      color: var(--text);
      margin: 0; padding: 0;
    }
    /* Nav */
    .app-bar {
      background-color: var(--surface);
      padding: 0 24px;
      height: 64px;
      display: flex; align-items: center; justify-content: space-between;
      box-shadow: 0 2px 4px rgba(0,0,0,0.3);
      position: sticky; top: 0; z-index: 100;
    }
    .app-bar-title { font-size: 1.25rem; font-weight: 500; color: var(--accent); }
    .nav-links a { color: var(--text-sec); text-decoration: none; margin-left: 20px; font-size: 0.9rem; transition: 0.2s; }
    .nav-links a:hover, .nav-links a.active { color: var(--accent); }

    /* Layout */
    .main { max-width: 1200px; margin: 24px auto; padding: 0 16px; display: grid; grid-template-columns: 1fr 380px; gap: 24px; }
    @media (max-width: 900px) { .main { grid-template-columns: 1fr; } }
    
    .card {
      background: var(--surface); border-radius: 8px; padding: 24px;
      box-shadow: 0 1px 3px rgba(0,0,0,0.2); margin-bottom: 24px;
    }
    h3 { font-size: 1.1rem; font-weight: 500; margin: 0 0 16px 0; color: var(--text); border-bottom: 1px solid var(--border); padding-bottom: 8px; }
    
    /* Inputs */
    .field { margin-bottom: 16px; }
    label { display: block; font-size: 0.75rem; font-weight: 500; color: var(--text-sec); margin-bottom: 6px; letter-spacing: 0.5px; text-transform: uppercase; }
    input[type="text"], textarea {
      width: 100%; background: #121212; border: 1px solid var(--border);
      color: var(--text); padding: 12px; border-radius: 4px; font-family: 'Roboto', sans-serif; font-size: 0.9rem;
      transition: border-color 0.2s;
    }
    input[type="text"]:focus, textarea:focus { outline: none; border-color: var(--accent); }
    textarea { min-height: 150px; resize: vertical; }

    /* Button */
    .btn {
      background: linear-gradient(135deg, #7C4DFF, #448AFF);
      color: white; border: none; padding: 12px 24px; border-radius: 4px;
      font-size: 0.95rem; font-weight: 500; letter-spacing: 0.5px; cursor: pointer;
      width: 100%; text-transform: uppercase; box-shadow: 0 4px 6px rgba(0,0,0,0.3);
      transition: transform 0.1s, box-shadow 0.2s;
    }
    .btn:hover { box-shadow: 0 6px 12px rgba(0,0,0,0.4); }
    .btn:active { transform: translateY(1px); }
    .btn.small { width: auto; padding: 6px 12px; font-size: 0.75rem; background: var(--surface-2); border: 1px solid var(--border); box-shadow: none; margin-top: 8px; }
    .btn.small:hover { background: #333; }

    /* Team Config (Right Col) */
    .team-card {
      background: var(--surface-2); border-radius: 6px; padding: 12px; margin-bottom: 12px; border: 1px solid var(--border);
    }
    .team-name { font-weight: 700; color: #fff; margin-bottom: 8px; display: flex; justify-content: space-between; font-size: 0.95rem; }
    .skill-row { display: flex; align-items: center; margin-bottom: 8px; font-size: 0.8rem; }
    .skill-lbl { width: 80px; overflow: hidden; white-space: nowrap; text-overflow: ellipsis; color: var(--text-sec); }
    .skill-val { width: 30px; text-align: right; margin-right: 8px; font-family: 'Roboto Mono', monospace; }
    input[type=range] { flex: 1; margin: 0 8px; accent-color: var(--accent); cursor: pointer; }
    
    /* Results */
    #results { margin-top: 24px; }
    .res-section { margin-bottom: 20px; }
    .res-label { color: var(--text-sec); font-size: 0.8rem; margin-bottom: 4px; }
    .res-val { font-size: 1rem; color: #fff; }
    
    .priority-badge {
        display: inline-block; padding: 4px 12px; border-radius: 4px; font-weight: bold; text-transform: uppercase; font-size: 0.8rem;
    }
    .p-Critical { background: rgba(244, 67, 54, 0.2); color: #ef5350; border: 1px solid #ef5350; }
    .p-High { background: rgba(255, 167, 38, 0.2); color: #ffa726; border: 1px solid #ffa726; }
    .p-Medium { background: rgba(102, 187, 106, 0.2); color: #66bb6a; border: 1px solid #66bb6a; }
    .p-Low { background: rgba(41, 182, 246, 0.2); color: #29b6f6; border: 1px solid #29b6f6; }
    
    .cat-chip {
        display: inline-block; background: #333; padding: 4px 10px; border-radius: 12px; font-size: 0.8rem; margin-right: 6px; border: 1px solid #444;
    }
    
    .csuite-box { background: rgba(187, 134, 252, 0.08); border-left: 3px solid var(--accent); padding: 8px 12px; font-size: 0.9rem; margin-top: 4px; }
    
    pre { background: #000; padding: 12px; border-radius: 4px; font-family: 'Roboto Mono', monospace; font-size: 0.8rem; color: #ccc; overflow-x: auto; border: 1px solid #333; }
    
    .json-area { font-family: 'Roboto Mono'; font-size: 0.75rem; min-height: 80px; }
    
    /* File input hidden */
    input[type="file"] { display: none; }
  </style>
</head>
<body>
  <div class="app-bar">
    <div class="app-bar-title">Intake</div>
    <div class="nav-links">
      <a href="/ui">Laws</a>
      <a href="/ui/intake" class="active">Intake</a>
      <a href="/ui/contracts">Contracts</a>
    </div>
  </div>

  <div class="main">
    <div>
       <div class="card">
         <h3>Inbound Analysis</h3>
         
         <div class="field">
           <label>Organization (Optional)</label>
           <input type="text" id="org" placeholder="e.g. Acme Corp">
         </div>
         
         <div class="field">
           <label>Inbound Message / Email Body</label>
           <textarea id="email_text" placeholder="Paste the full email or request here..."></textarea>
         </div>
         
         <div class="field">
            <label>Notify Email (Optional)</label>
            <input type="text" id="notify_email" placeholder="result@example.com">
         </div>
         
         <button id="analyze_btn" class="btn">Analyze & Route</button>
         <div id="status" style="margin-top:12px; font-size:0.9rem; color: var(--text-sec); font-style: italic;"></div>
       </div>

       <div id="results"></div>
       
       <div class="card" style="margin-top:24px;">
         <h3>Configuration Context</h3>
         <div class="field">
            <label>Playbook / Reference Notes (Overrides Team Routing if specific)</label>
            <textarea id="ref_notes" style="min-height:80px; font-size:0.85rem;"></textarea>
         </div>
         <div class="field">
            <label>Watchlist (C-Suite)</label>
            <input type="text" id="csuite" placeholder="e.g. CEO, Jane Doe">
         </div>
       </div>
    </div>

    <div>
      <div class="card">
        <h3>Team Routing Profile</h3>
        <p style="font-size:0.8rem; color:var(--text-sec); margin-bottom:16px;">
           Configure team skills to simulate routing logic. Adjust mastery (0-100).
        </p>
        
        <div id="team_profile_container"></div>
        
        <button class="btn small" id="btn_add_member">+ Add Member</button>
        
        <div style="margin-top:24px; border-top:1px solid var(--border); padding-top:16px;">
            <label>JSON Profile Import/Export</label>
            <textarea id="team_json" class="json-area"></textarea>
            
            <div style="display:flex; gap:8px; flex-wrap:wrap;">
                <button class="btn small" id="btn_team_apply">Apply Text</button>
                <button class="btn small" id="btn_team_reset">Reset</button>
                <label for="import_file" class="btn small" style="cursor:pointer; display:inline-block; text-align:center;">
                   Import JSON File
                </label>
                <input type="file" id="import_file" accept=".json">
            </div>
        </div>
      </div>
    </div>
  </div>

  <script>
    // --- Defaults ---
    const defaultNotes = "# Playbook & Routing Rules\\n\\n## 1. Commercial & Contracts\\n- Keywords: contract, agreement, sow, msa, nda, negotiation, renewal\\n- Primary Owner: Ron\\n- Priority: Medium (unless 'urgent' or 'today' mentioned)\\n\\n## 2. Privacy & Cybersecurity\\n- Keywords: breach, incident, gdpr, ccpa, dpa, security, privacy\\n- Primary Owner: Shawn\\n- Priority: High (Critical if 'breach' or 'incident')\\n\\n## 3. Data & Compliance\\n- Keywords: analytics, ai, data usage, compliance, audit, tax\\n- Primary Owner: Doug\\n- Priority: Medium\\n\\n## 4. Litigation & Disputes\\n- Keywords: lawsuit, subpoena, court, dispute, cease and desist\\n- Primary Owner: Ron\\n- Priority: High";
    
    const defaultTeamProfile = {
      members: [
        {
          name: "Shawn",
          skills: [
            { label: "saas", mastery: 95 },
            { label: "cybersecurity", mastery: 95 },
            { label: "privacy", mastery: 95 },
            { label: "contracts", mastery: 100 },
            { label: "ai", mastery: 100 },
            { label: "real estate", mastery: 55 }           
          ]
        },
        {
          name: "Ron",
          skills: [
            { label: "negotiating", mastery: 95 },
            { label: "contracts", mastery: 95 },
            { label: "litigation", mastery: 95 },
            { label: "real estate", mastery: 95 }
          ]
        },
        {
          name: "Russell",
          skills: [
            { label: "open source", mastery: 95 },
            { label: "compliance", mastery: 95 },
            { label: "litigation", mastery: 85 }
          ]
        }
      ]
    };

    let teamProfile = JSON.parse(JSON.stringify(defaultTeamProfile));

    // --- Init ---
    document.addEventListener("DOMContentLoaded", () => {
      const ref = document.getElementById("ref_notes");
      if (ref && !ref.value.trim()) ref.value = defaultNotes;
      
      renderTeamProfile();
      
      document.getElementById("btn_team_reset").addEventListener("click", () => {
         teamProfile = JSON.parse(JSON.stringify(defaultTeamProfile));
         renderTeamProfile();
      });
      document.getElementById("btn_team_apply").addEventListener("click", applyTeamFromJson);
      document.getElementById("btn_add_member").addEventListener("click", addMember);
      document.getElementById("analyze_btn").addEventListener("click", analyzeIntake);
      
      // Import File Listener
      document.getElementById("import_file").addEventListener("change", handleFileImport);
    });

    // --- Team UI Render Logic ---
    function renderTeamProfile() {
      const container = document.getElementById("team_profile_container");
      container.innerHTML = "";
      
      teamProfile.members.forEach((member, mi) => {
        const card = document.createElement("div");
        card.className = "team-card";
        
        const header = document.createElement("div");
        header.className = "team-name";
        header.innerHTML = `<span>${member.name}</span> <span style='cursor:pointer; opacity:0.5;' onclick='removeMember(${mi})'>&times;</span>`;
        card.appendChild(header);
        
        member.skills.forEach((skill, si) => {
           const row = document.createElement("div");
           row.className = "skill-row";
           
           // Label
           const lbl = document.createElement("div");
           lbl.className = "skill-lbl";
           lbl.textContent = skill.label;
           lbl.title = skill.label;
           
           // Slider
           const slider = document.createElement("input");
           slider.type = "range"; slider.min=0; slider.max=100;
           slider.value = skill.mastery;
           slider.oninput = (e) => {
              teamProfile.members[mi].skills[si].mastery = parseInt(e.target.value);
              valDisp.textContent = e.target.value;
              syncJson();
           };
           
           // Value
           const valDisp = document.createElement("div");
           valDisp.className = "skill-val";
           valDisp.textContent = skill.mastery;
           
           row.appendChild(lbl);
           row.appendChild(slider);
           row.appendChild(valDisp);
           card.appendChild(row);
        });
        
        const addSkillBtn = document.createElement("div");
        addSkillBtn.style.textAlign = "center";
        addSkillBtn.innerHTML = "<span style='font-size:0.7rem; color:#666; cursor:pointer;'>+ Add Skill</span>";
        addSkillBtn.onclick = () => addSkill(mi);
        card.appendChild(addSkillBtn);
        
        container.appendChild(card);
      });
      syncJson();
    }
    
    // --- Team Logic ---
    
    function addSkill(mi) {
        const lbl = prompt("Skill Name (e.g. litigation)");
        if(lbl) {
            teamProfile.members[mi].skills.push({ label: lbl, mastery: 50 });
            renderTeamProfile();
        }
    }
    
    function addMember() {
        const name = prompt("Member Name:");
        if(name) {
            teamProfile.members.push({ name: name, skills: [] });
            renderTeamProfile();
        }
    }
    
    // Global scope for HTML access
    window.removeMember = function(mi) {
        if(confirm("Remove this member?")) {
            teamProfile.members.splice(mi, 1);
            renderTeamProfile();
        }
    };
    
    function syncJson() {
        document.getElementById("team_json").value = JSON.stringify(teamProfile, null, 2);
    }
    
    function applyTeamFromJson() {
        try {
            const parsed = JSON.parse(document.getElementById("team_json").value);
            if(parsed && parsed.members) {
                teamProfile = parsed;
                renderTeamProfile();
            }
        } catch(e) { alert("Invalid JSON"); }
    }
    
    function handleFileImport(e) {
        const file = e.target.files[0];
        if (!file) return;
        
        const reader = new FileReader();
        reader.onload = function(evt) {
            try {
                const parsed = JSON.parse(evt.target.result);
                if (parsed && parsed.members) {
                    teamProfile = parsed;
                    renderTeamProfile();
                    alert("Team profile imported successfully.");
                } else {
                    alert("Invalid JSON format. Must contain 'members' array.");
                }
            } catch(err) {
                alert("Error parsing JSON file: " + err);
            }
        };
        reader.readAsText(file);
        // Reset input so same file can be selected again if needed
        e.target.value = '';
    }

    // --- Intake Analysis Logic ---
    async function analyzeIntake() {
       const btn = document.getElementById("analyze_btn");
       const status = document.getElementById("status");
       const resultsDiv = document.getElementById("results");
       
       const email = document.getElementById("email_text").value;
       if(!email.trim()) { alert("Please enter message text."); return; }
       
       btn.disabled = true;
       status.textContent = " Analyzing content & routing...";
       resultsDiv.innerHTML = "";
       
       const payload = {
           email_text: email,
           reference_notes: document.getElementById("ref_notes").value,
           organization_name: document.getElementById("org").value,
           csuite_names: document.getElementById("csuite").value.split(",").map(s=>s.trim()).filter(s=>s),
           max_categories: 5, // Hardcoded default since UI element removed
           notify_email: document.getElementById("notify_email").value,
           team_profile: teamProfile
       };
       
       try {
           const res = await fetch("/api/intake/analyze", {
               method: "POST",
               headers: {"Content-Type": "application/json"},
               body: JSON.stringify(payload)
           });
           const data = await res.json();
           
           status.textContent = " Analysis complete.";
           
           // Render Results Card
           const card = document.createElement("div");
           card.className = "card";
           card.style.borderTop = "4px solid var(--accent)";
           
           // Priority Header
           const pLabel = data.priority_label || "Normal";
           const pClass = "p-" + pLabel;
           
           let html = `
             <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:16px;">
                <h3>Analysis Result</h3>
                <span class="priority-badge ${pClass}">${pLabel} (${data.priority_score}/10)</span>
             </div>
           `;
           
           // Categories
           html += `<div class="res-section"><div class="res-label">CATEGORIES</div>`;
           if(data.categories && data.categories.length) {
               data.categories.forEach(c => { html += `<span class="cat-chip">${c}</span>`; });
           } else { html += `<span style="color:#666">None</span>`; }
           html += `</div>`;
           
           // Summary
           html += `<div class="res-section"><div class="res-label">SUMMARY</div><div class="res-val" style="line-height:1.4">${data.summary}</div></div>`;
           
           // C-Suite
           if(data.csuite_mentions && data.csuite_mentions.length > 0) {
              html += `<div class="res-section"><div class="res-label">EXECUTIVE MENTIONS</div>`;
              data.csuite_mentions.forEach(m => {
                  html += `<div class="csuite-box"><strong>${m.name}</strong> detected.</div>`;
              });
              html += `</div>`;
           }
           
           // Routing
           html += `<div class="res-section" style="background:#222; padding:12px; border-radius:6px; border:1px solid #333;">
              <div class="res-label" style="color:var(--primary);">SUGGESTED OWNER</div>
              <div style="font-size:1.1rem; font-weight:bold; color:#fff;">${data.suggested_owner || 'Unassigned'}</div>`;
              
           if(data.suggested_backup) {
              html += `<div style="font-size:0.85rem; color:#aaa; margin-top:4px;">Backup: ${data.suggested_backup}</div>`;
           }
           if(data.learning_opportunities && data.learning_opportunities.length) {
              html += `<div style="font-size:0.85rem; color:var(--accent); margin-top:8px;">Suggested Training: ${data.learning_opportunities.join(", ")}</div>`;
           }
           html += `</div>`;
           
           // Next Steps
           if(data.suggested_next_steps) {
               html += `<div class="res-section"><div class="res-label">NEXT STEPS</div><pre style="white-space:pre-wrap; background:#1a1a1a;">${data.suggested_next_steps}</pre></div>`;
           }
           
           // Email Status
           if(data.email_status) {
              html += `<div style="font-size:0.75rem; color:#666; text-align:right;">Email Notification: ${data.email_status}</div>`;
           }
           
           // --- NEW: Original Request Text Block ---
           if (data.original_text) {
               html += `<div class="res-section" style="margin-top:20px; padding-top:16px; border-top:1px solid #333;">
                  <div class="res-label" style="margin-bottom:8px;">ORIGINAL REQUEST</div>
                  <div style="font-size:0.85rem; color:#bbb; white-space:pre-wrap; background:#111; padding:12px; border-radius:4px; font-family:'Roboto Mono', monospace;">${data.original_text}</div>
               </div>`;
           }

           card.innerHTML = html;
           resultsDiv.appendChild(card);
           
       } catch(e) {
           console.error(e);
           status.textContent = "Error: " + e;
       } finally {
           btn.disabled = false;
       }
    }
  </script>
</body>
</html>
    """
    return HTMLResponse(content=html)

# ---------- Contract Redline UI at /ui/contracts ----------

@app.get("/ui/contracts", response_class=HTMLResponse)
async def ui_contracts():
    html = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Phoenix Contract Redline</title>
  <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&family=Roboto+Mono:wght@400;500&display=swap" rel="stylesheet">
  <style>
    :root { --primary: #90CAF9; --bg: #121212; --surface: #1E1E1E; --surface-2: #2C2C2C; --text: #E0E0E0; --text-sec: #A0A0A0; --border: #333; --accent: #FFA500; }
    * { box-sizing: border-box; }
    body { font-family: 'Roboto', sans-serif; background-color: var(--bg); color: var(--text); margin: 0; padding: 0; }
    
    .app-bar { background-color: var(--surface); padding: 0 24px; height: 64px; display: flex; align-items: center; justify-content: space-between; box-shadow: 0 2px 4px rgba(0,0,0,0.3); position: sticky; top: 0; z-index: 100; }
    .app-bar-title { font-size: 1.25rem; font-weight: 500; color: var(--accent); }
    .nav-links a { color: var(--text-sec); text-decoration: none; margin-left: 20px; font-size: 0.9rem; transition: 0.2s; }
    .nav-links a.active { color: var(--accent); }

    .main { max-width: 1100px; margin: 24px auto; padding: 0 16px; display: grid; grid-template-columns: 2fr 1fr; gap: 24px; }
    @media (max-width: 900px) { .main { grid-template-columns: 1fr; } }

    .card { background: var(--surface); border-radius: 8px; padding: 24px; box-shadow: 0 1px 3px rgba(0,0,0,0.2); margin-bottom: 24px; }
    h2 { font-weight: 400; font-size: 1.3rem; margin: 0 0 16px 0; color: var(--text); border-bottom: 1px solid var(--border); padding-bottom: 8px; }
    
    .field { margin-bottom: 16px; }
    label { display: block; font-size: 0.8rem; font-weight: 500; color: var(--text-sec); margin-bottom: 8px; text-transform: uppercase; }
    select, input[type=file], input[type=text] { width: 100%; padding: 10px; background: var(--surface-2); border: 1px solid var(--border); color: var(--text); border-radius: 4px; font-size: 0.9rem; }
    textarea { width: 100%; padding: 10px; background: var(--surface-2); border: 1px solid var(--border); color: var(--text); border-radius: 4px; font-size: 0.9rem; min-height: 150px; font-family: 'Roboto Mono', monospace; }
    
    .btn { background: var(--accent); color: #000; border: none; padding: 12px 24px; border-radius: 4px; font-size: 0.95rem; font-weight: 500; cursor: pointer; width: 100%; text-transform: uppercase; box-shadow: 0 2px 4px rgba(0,0,0,0.2); }
    .btn:hover { filter: brightness(1.1); }
    .btn:disabled { opacity: 0.6; cursor: not-allowed; }
    .btn-small { padding: 6px 12px; font-size: 0.8rem; width: auto; }
    .btn-danger { background: #ef5350; color: white; }

    /* Analysis Results */
    .analysis-item { background: #252525; border-left: 4px solid #555; margin-bottom: 15px; padding: 15px; border-radius: 4px; }
    .analysis-item.has-changes { border-left-color: #ef5350; }
    .clause-label { font-size: 0.75rem; color: #888; text-transform: uppercase; font-weight: bold; margin-bottom: 6px; }
    .clause-box { font-family: 'Roboto Mono', monospace; font-size: 0.8rem; background: #111; padding: 10px; border-radius: 4px; color: #ccc; white-space: pre-wrap; max-height: 200px; overflow-y: auto; border: 1px solid #333; }
    .change-row { margin-bottom: 8px; font-family: 'Roboto Mono', monospace; font-size: 0.85rem; display: flex; align-items: flex-start; margin-top: 8px; }
    .badge { display: inline-block; padding: 2px 6px; border-radius: 3px; font-size: 0.7em; font-weight: bold; margin-right: 10px; min-width: 70px; text-align: center; flex-shrink: 0; }
    .badge.ins { background: rgba(102, 187, 106, 0.15); color: #66bb6a; border: 1px solid #66bb6a; }
    .badge.del { background: rgba(239, 83, 80, 0.15); color: #ef5350; border: 1px solid #ef5350; }
    .badge.rep { background: rgba(255, 167, 38, 0.15); color: #ffa726; border: 1px solid #ffa726; }
    .badge.cmt { background: rgba(66, 165, 245, 0.15); color: #42a5f5; border: 1px solid #42a5f5; }
    
    /* Persona List in Sidebar */
    .persona-list-item { padding: 10px; background: var(--surface-2); margin-bottom: 8px; border-radius: 4px; cursor: pointer; border: 1px solid transparent; display: flex; justify-content: space-between; align-items: center; }
    .persona-list-item:hover { border-color: var(--primary); }
    .persona-list-item.active { background: rgba(144, 202, 249, 0.15); border-color: var(--primary); color: var(--primary); }
  </style>
</head>
<body>
  <div class="app-bar">
    <div class="app-bar-title">Contract Redline</div>
    <div class="nav-links">
      <a href="/ui">Laws</a>
      <a href="/ui/intake">Intake</a>
      <a href="/ui/contracts" class="active">Contracts</a>
    </div>
  </div>

  <div class="main">
    
    <div>
      <div class="card">
        <h2>1. Upload & Configure</h2>
        
        <div class="field">
          <label>Counterparty DOCX (Required)</label>
          <input id="counterparty" type="file" accept=".docx">
        </div>
        <div class="field">
          <label>Template DOCX (Optional)</label>
          <input id="template" type="file" accept=".docx">
        </div>

        <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 16px;">
            <div class="field">
              <label>Our Role</label>
              <select id="role_select">
                  <option value="Buyer">Buyer / Client</option>
                  <option value="Seller">Seller / Provider</option>
              </select>
            </div>
            <div class="field">
              <label>Persona</label>
              <select id="persona_select">
                  <option value="General Counsel">General Counsel</option>
                  </select>
            </div>
        </div>

        <button class="btn" id="analyze_btn" onclick="runRedline()">Analyze Contract</button>
      </div>

      <div class="card" style="border-top: 4px solid var(--accent); min-height:100px;">
        <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom: 20px;">
            <h2 style="margin:0; border:none;">2. Analysis Result</h2>
            <div id="export_area" style="display:none; display:flex; gap:12px;">
                <button class="btn btn-small" style="background:#66BB6A;" onclick="downloadRedline()">Download Redline</button>
                <button class="btn btn-small" style="background:#42A5F5;" onclick="downloadReport()">Download Report</button>
            </div>
        </div>
        <div id="status_text" style="color:#777; font-style:italic;">Ready to analyze...</div>
        <div id="output"></div>
      </div>
    </div>

    <div>
      <div class="card">
        <h2>Persona Library</h2>
        <div id="persona_list"></div>
        <button class="btn btn-small" style="margin-top:12px; width:100%; background:#444;" onclick="createNewPersona()">+ New Persona</button>
      </div>

      <div class="card" id="editor_card" style="display:none;">
        <h2 style="font-size:1rem;">Edit Persona</h2>
        <div class="field">
            <label>Name</label>
            <input type="text" id="edit_name">
        </div>
        <div class="field">
            <label>Instructions / Strategy</label>
            <textarea id="edit_instructions"></textarea>
        </div>
        <div style="display:flex; gap:8px;">
            <button class="btn btn-small" onclick="savePersona()">Save</button>
            <button class="btn btn-small btn-danger" onclick="deletePersona()">Delete</button>
        </div>
      </div>
    </div>

  </div>

  <script>
    let lastAnalysisData = null;
    let lastUploadedFile = null;
    let availablePersonas = [];

    // --- Init ---
    document.addEventListener("DOMContentLoaded", () => {
        loadPersonas();
    });

    async function loadPersonas() {
        try {
            const res = await fetch("/api/contracts/personas");
            const data = await res.json();
            availablePersonas = data;
            renderPersonaList();
            renderPersonaSelect();
        } catch(e) { console.error("Failed to load personas", e); }
    }

    function renderPersonaList() {
        const list = document.getElementById("persona_list");
        list.innerHTML = "";
        availablePersonas.forEach(p => {
            const div = document.createElement("div");
            div.className = "persona-list-item";
            div.innerText = p.name;
            div.onclick = () => openEditor(p);
            list.appendChild(div);
        });
    }

    function renderPersonaSelect() {
        const sel = document.getElementById("persona_select");
        const currentVal = sel.value; 
        sel.innerHTML = "";
        availablePersonas.forEach(p => {
            const opt = document.createElement("option");
            opt.value = p.name;
            opt.innerText = p.name;
            sel.appendChild(opt);
        });
        if (availablePersonas.find(p => p.name === currentVal)) {
            sel.value = currentVal;
        }
    }

    // --- Editor Logic ---
    function openEditor(persona) {
        document.getElementById("editor_card").style.display = "block";
        document.getElementById("edit_name").value = persona.name;
        document.getElementById("edit_name").disabled = true; // Edit existing by name lock
        document.getElementById("edit_instructions").value = persona.instructions;
    }

    function createNewPersona() {
        document.getElementById("editor_card").style.display = "block";
        document.getElementById("edit_name").value = "";
        document.getElementById("edit_name").disabled = false;
        document.getElementById("edit_instructions").value = "Tone: ...\\nStrategy: ...";
    }

    async function savePersona() {
        const name = document.getElementById("edit_name").value;
        const instr = document.getElementById("edit_instructions").value;
        if(!name) return alert("Name required");

        await fetch("/api/contracts/personas", {
            method: "POST",
            headers: {"Content-Type": "application/json"},
            body: JSON.stringify({ name: name, instructions: instr })
        });
        loadPersonas();
        alert("Saved");
    }

    async function deletePersona() {
        const name = document.getElementById("edit_name").value;
        if(!confirm("Delete " + name + "?")) return;
        
        await fetch("/api/contracts/personas/" + encodeURIComponent(name), { method: "DELETE" });
        document.getElementById("editor_card").style.display = "none";
        loadPersonas();
    }

    // --- Redline Logic ---
    
    async function runRedline() {
        const output = document.getElementById("output");
        const status = document.getElementById("status_text");
        const btn = document.getElementById("analyze_btn");
        const cpFile = document.getElementById("counterparty").files[0];
        
        if (!cpFile) return alert("Upload a counterparty DOCX!");

        output.innerHTML = "";
        document.getElementById("export_area").style.display = "none";
        status.innerText = "Uploading & Analyzing...";
        btn.disabled = true;
        lastUploadedFile = cpFile;

        const formData = new FormData();
        formData.append("counterparty", cpFile);
        const tpFile = document.getElementById("template").files[0];
        if (tpFile) formData.append("template", tpFile);

        try {
            // 1. Upload
            const upRes = await fetch("/api/contracts/redline/upload", { method: "POST", body: formData });
            if (!upRes.ok) throw new Error(await upRes.text());
            const upData = await upRes.json();

            // 2. Analyze
            const payload = {
                counterparty_text: upData.counterparty_text,
                template_text: upData.template_text,
                mode: "template_only", 
                persona: document.getElementById("persona_select").value,
                role: document.getElementById("role_select").value // NEW
            };

            const anRes = await fetch("/api/contracts/redline/analyze", {
                method: "POST",
                headers: { "Content-Type": "application/json" },
                body: JSON.stringify(payload)
            });
            
            if (!anRes.ok) throw new Error(await anRes.text());
            const data = await anRes.json();
            
            lastAnalysisData = data.diff;
            status.innerText = `Analysis Complete. ${data.diff.length} issues found.`;
            status.style.color = "#66BB6A";
            
            // Show Export Area with Flex for buttons
            const exp = document.getElementById("export_area");
            exp.style.display = "flex";
            exp.style.gap = "12px";
            
            renderResults(data.diff);

        } catch(e) {
            status.innerText = "Error: " + e.message;
            status.style.color = "#ef5350";
        } finally {
            btn.disabled = false;
        }
    }

    function renderResults(diffs) {
        const container = document.getElementById("output");
        if (!diffs || diffs.length === 0) return container.innerHTML = "<div style='padding:20px; text-align:center;'>No redlines or issues found (Contract looks clean).</div>";

        diffs.forEach(item => {
            const d = item.delta || {};
            const ins = d.insertions || [];
            const del = d.deletions || [];
            const rep = d.replacements || [];
            const cmt = d.comments || [];
            const hasChanges = (ins.length + del.length + rep.length + cmt.length) > 0;

            const card = document.createElement("div");
            card.className = "analysis-item" + (hasChanges ? " has-changes" : "");
            
            let html = `<div><div class="clause-label">Clause</div><div class="clause-box">${item.cp_text || "(New Clause)"}</div></div>`;
            
            if (hasChanges) {
                ins.forEach(x => html += `<div class="change-row"><span class="badge ins">INSERT</span><span>${x}</span></div>`);
                del.forEach(x => html += `<div class="change-row"><span class="badge del">DELETE</span><span>${x}</span></div>`);
                rep.forEach(x => html += `<div class="change-row"><span class="badge rep">REPLACE</span><span>"${x.from}" &rarr; "${x.to}"</span></div>`);
                cmt.forEach(x => html += `<div class="change-row"><span class="badge cmt">NOTE</span><span>${x}</span></div>`);
            } else {
                html += `<div style="margin-top:8px; font-size:0.8rem; color:#666;">No issues found.</div>`;
            }
            card.innerHTML = html;
            container.appendChild(card);
        });
    }

    async function downloadRedline() {
        if (!lastAnalysisData) return;
        const toBase64 = file => new Promise((resolve, reject) => {
            const r = new FileReader(); r.readAsDataURL(file);
            r.onload = () => resolve(r.result.split(',')[1]); r.onerror = reject;
        });
        
        try {
            const b64 = await toBase64(lastUploadedFile);
            const res = await fetch("/api/contracts/redline/export", {
                method: "POST",
                headers: { "Content-Type": "application/json" },
                body: JSON.stringify({ original_docx_base64: b64, diff: lastAnalysisData })
            });
            if(!res.ok) throw new Error(await res.text());
            
            const blob = await res.blob();
            const url = window.URL.createObjectURL(blob);
            const a = document.createElement("a"); a.href = url; a.download = "redlined.docx";
            document.body.appendChild(a); a.click(); document.body.removeChild(a);
        } catch(e) { alert(e); }
    }
    
    async function downloadReport() {
        if (!lastAnalysisData) return;
        try {
            const res = await fetch("/api/contracts/redline/export-report", {
                method: "POST",
                headers: { "Content-Type": "application/json" },
                body: JSON.stringify({ diff: lastAnalysisData })
            });
            if(!res.ok) throw new Error(await res.text());
            
            const blob = await res.blob();
            const url = window.URL.createObjectURL(blob);
            const a = document.createElement("a"); a.href = url; a.download = "Analysis_Report.docx";
            document.body.appendChild(a); a.click(); document.body.removeChild(a);
        } catch(e) { alert(e); }
    }
  </script>
</body>
</html>
    """
    return HTMLResponse(content=html)