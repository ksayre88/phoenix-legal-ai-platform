from fastapi import APIRouter, UploadFile, File, HTTPException, Body
from fastapi.responses import StreamingResponse
from typing import Optional
import base64
import io
from docx import Document
from docx.shared import RGBColor

# Import Schemas
from app.models.schemas import (
    ContractRedlineRequest, 
    PersonaUpdateRequest, 
    ContractRedlineExportRequest, 
    ContractReportRequest
)

# Import the logic from the Service layer
from app.services.contracts import (
    analyze_contract_logic, 
    get_personas, 
    upsert_persona, 
    delete_persona
)

from app.utils.file_parsing import extract_docx_text
from app.utils.redline_apply import apply_redlines_to_docx

router = APIRouter()

# --- Persona Management ---
@router.get("/personas")
async def get_contract_personas():
    return get_personas()

@router.post("/personas")
async def upsert_contract_persona(req: PersonaUpdateRequest):
    upsert_persona(req.name, req.instructions)
    return {"status": "ok"}

@router.delete("/personas/{name}")
async def delete_contract_persona(name: str):
    delete_persona(name)
    return {"status": "ok"}

# --- Document Handling ---
@router.post("/redline/upload")
async def upload_contracts(counterparty: UploadFile = File(...), template: UploadFile = File(None)):
    """
    Parses .docx files into text for the frontend editor.
    """
    try:
        cp_bytes = await counterparty.read()
        tp_text = None
        if template:
            tp_bytes = await template.read()
            tp_text = extract_docx_text(tp_bytes)
        
        return {
            "status": "ok", 
            "counterparty_text": extract_docx_text(cp_bytes), 
            "template_text": tp_text
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"File parsing error: {e}")

@router.post("/redline/analyze")
async def analyze_route(req: ContractRedlineRequest):
    """
    Triggers the AI analysis.
    """
    try:
        # Calls the function in services/contracts.py
        result = await analyze_contract_logic(
            counterparty_text=req.counterparty_text,
            template_text=req.template_text,
            persona=req.persona
        )
        return result
    except Exception as e:
        # Print error to console for debugging
        print(f"Analysis Failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/redline/export")
async def export_redline(req: ContractRedlineExportRequest):
    """
    Generates the actual .docx file with Track Changes applied.
    """
    try:
        orig = base64.b64decode(req.original_docx_base64)
        res = apply_redlines_to_docx(orig, req.diff)
        
        return StreamingResponse(
            io.BytesIO(res), 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={"Content-Disposition": "attachment; filename=redlined.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

@router.post("/redline/export-report")
async def export_report(req: ContractReportRequest):
    """
    Simple summary report generation.
    """
    try:
        doc = Document()
        doc.add_heading('Phoenix Analysis Report', 0)
        
        for i, item in enumerate(req.diff, 1):
            # Extract data safely
            clause_name = item.get("clause_name", f"Clause {i}")
            cp_text = item.get("original_text", "") or item.get("cp_text", "")
            
            # Handle flattened or nested delta structure
            delta = item.get("delta", item)
            risk_score = item.get("risk_score") or delta.get("risk_score", "N/A")

            # 1. Heading
            doc.add_heading(clause_name, level=2)
            
            # 2. Original Text (Italic)
            p_text = doc.add_paragraph()
            run_text = p_text.add_run(f"Original Text: \"{cp_text[:200]}...\"")
            run_text.italic = True
            
            # 3. Risk Score (Bold) - FIXED LINE
            # Instead of style='Strong', we use a normal paragraph and bold the run manually
            p_risk = doc.add_paragraph()
            run_risk = p_risk.add_run(f"Risk Score: {risk_score}/10")
            run_risk.bold = True
            run_risk.font.color.rgb = RGBColor(200, 0, 0) if str(risk_score) > "5" else RGBColor(0, 0, 0)
                 
            # 4. Comments / Reasoning
            reasoning = delta.get("reasoning", "")
            if reasoning:
                doc.add_paragraph(f"Reasoning: {reasoning}")

            if delta.get("comments"):
                for c in delta["comments"]: 
                    doc.add_paragraph(c, style='List Bullet')
            
            doc.add_paragraph("_" * 50) # Separator

        buf = io.BytesIO()
        doc.save(buf)
        return StreamingResponse(
            io.BytesIO(buf.getvalue()), 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Analysis_Report.docx"}
        )
    except Exception as e:
        print(f"Report Generation Failed: {e}")
        raise HTTPException(status_code=500, detail=f"Report failed: {str(e)}")